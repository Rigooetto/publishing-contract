from flask import Flask, render_template_string, request, send_file, session, redirect, url_for, jsonify, flash
from docx import Document
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy import func, or_
from babel.dates import format_date
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload
import io
import os
import re
import json
import zipfile
import datetime
import base64
import xml.etree.ElementTree as ET
import traceback

from docusign_esign import ApiClient, EnvelopesApi, EnvelopeDefinition, Document as DocusignDocument, Signer, SignHere, Tabs, Recipients

DOCUSIGN_ACCOUNT_ID = os.getenv("DOCUSIGN_ACCOUNT_ID", "")
DOCUSIGN_BASE_PATH = os.getenv("DOCUSIGN_BASE_PATH", "https://demo.docusign.net/restapi")
DOCUSIGN_AUTH_SERVER = os.getenv("DOCUSIGN_AUTH_SERVER", "account-d.docusign.com")
DOCUSIGN_INTEGRATION_KEY = os.getenv("DOCUSIGN_INTEGRATION_KEY", "")
DOCUSIGN_USER_ID = os.getenv("DOCUSIGN_USER_ID", "")
DOCUSIGN_PRIVATE_KEY = os.getenv("DOCUSIGN_PRIVATE_KEY", "")

app = Flask(__name__)
app.secret_key = os.getenv("SECRET_KEY", "change-this-secret-key")

raw_db_url = os.getenv("DATABASE_URL", "sqlite:///writers.db")
if raw_db_url.startswith("postgres://"):
    raw_db_url = raw_db_url.replace("postgres://", "postgresql://", 1)
if raw_db_url.startswith("postgresql://") and "sslmode=" not in raw_db_url:
    joiner = "&" if "?" in raw_db_url else "?"
    raw_db_url = raw_db_url + joiner + "sslmode=require"

app.config["SQLALCHEMY_DATABASE_URI"] = raw_db_url
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
app.config["MAX_CONTENT_LENGTH"] = 25 * 1024 * 1024

db = SQLAlchemy(app)

BASE_TEMPLATE_DIR = os.getenv("TEMPLATE_DIR", "template")
FULL_CONTRACT_TEMPLATE = os.getenv(
    "FULL_CONTRACT_TEMPLATE",
    os.path.join(BASE_TEMPLATE_DIR, "PUBLISHING_AGREEMENT_CONTRACT.docx"),
)
SCHEDULE_1_TEMPLATE = os.getenv(
    "SCHEDULE_1_TEMPLATE",
    os.path.join(BASE_TEMPLATE_DIR, "SCHEDULE_1.docx"),
)
OUTPUT_DIR = os.getenv("OUTPUT_DIR", "generated_contracts")

TEAM_USERNAME = os.getenv("TEAM_USERNAME")
TEAM_PASSWORD = os.getenv("TEAM_PASSWORD")

DEFAULT_PUBLISHER_ADDRESS = "3840 E. Miraloma Ave"
DEFAULT_PUBLISHER_CITY = "Anaheim"
DEFAULT_PUBLISHER_STATE = "CA"
DEFAULT_PUBLISHER_ZIP = "92806"

GOOGLE_DRIVE_FOLDER_ID = os.getenv("GOOGLE_DRIVE_FOLDER_ID", "")
GOOGLE_SERVICE_ACCOUNT_JSON = os.getenv("GOOGLE_SERVICE_ACCOUNT_JSON", "")

app.logger.warning("ENV CHECK: folder=%s json=%s", bool(GOOGLE_DRIVE_FOLDER_ID), bool(GOOGLE_SERVICE_ACCOUNT_JSON))
app.logger.warning("JSON LEN: %s", len(GOOGLE_SERVICE_ACCOUNT_JSON or ""))

os.makedirs(OUTPUT_DIR, exist_ok=True)


def slugify(value):
    value = (value or "").strip()
    value = re.sub(r"[^\w\s-]", "", value, flags=re.UNICODE)
    value = re.sub(r"[-\s]+", "_", value)
    return value or "file"


def parse_float(value):
    try:
        return float((value or "").strip())
    except ValueError:
        return 0.0


def build_full_name(first_name, middle_name, last_names):
    return " ".join(
        part.strip() for part in [first_name, middle_name, last_names] if part and part.strip()
    ).strip()


def normalize_text(value):
    return " ".join((value or "").lower().strip().split())


def normalize_title(title):
    return normalize_text(title)


def build_writer_identity_from_row(row):
    ipi = (row.get("ipi") or "").strip()
    if ipi:
        return "ipi:" + ipi.lower()
    selected_writer_id = (row.get("selected_writer_id") or "").strip()
    if selected_writer_id:
        return "id:" + selected_writer_id
    return "name:" + normalize_text(row.get("full_name", ""))


def build_writer_identity_from_workwriter(work_writer):
    if work_writer.writer and work_writer.writer.ipi:
        return "ipi:" + work_writer.writer.ipi.lower()
    if work_writer.writer_id:
        return "id:" + str(work_writer.writer_id)
    return "name:" + normalize_text(work_writer.writer.full_name if work_writer.writer else "")

def build_session_name(raw_name):
    raw_name = (raw_name or "").strip()
    prefix = datetime.datetime.utcnow().strftime("%Y.%m.%d")
    if raw_name:
        return prefix + " " + raw_name
    return prefix

def default_publisher_for_pro(pro):
    return {
        "BMI": "Songs of Afinarte",
        "ASCAP": "Melodies of Afinarte",
        "SESAC": "Music of Afinarte",
    }.get((pro or "").strip(), "")


def default_publisher_ipi_for_pro(pro):
    return {
        "BMI": "817874992",
        "ASCAP": "807953316",
        "SESAC": "817094629",
    }.get((pro or "").strip(), "")


def get_drive_service():
    if not GOOGLE_SERVICE_ACCOUNT_JSON:
        raise ValueError("Missing GOOGLE_SERVICE_ACCOUNT_JSON")
    info = json.loads(GOOGLE_SERVICE_ACCOUNT_JSON)
    credentials = service_account.Credentials.from_service_account_info(
        info, scopes=["https://www.googleapis.com/auth/drive"]
    )
    return build("drive", "v3", credentials=credentials, cache_discovery=False)


def get_docusign_api_client():
    private_key = (DOCUSIGN_PRIVATE_KEY or "").replace("\\n", "\n")
    private_key_bytes = private_key.encode("utf-8")
    api_client = ApiClient()
    api_client.host = DOCUSIGN_BASE_PATH
    token = api_client.request_jwt_user_token(
        client_id=DOCUSIGN_INTEGRATION_KEY,
        user_id=DOCUSIGN_USER_ID,
        oauth_host_name=DOCUSIGN_AUTH_SERVER,
        private_key_bytes=private_key_bytes,
        expires_in=3600,
        scopes=["signature", "impersonation"],
    )
    api_client.set_default_header("Authorization", "Bearer " + token.access_token)
    return api_client


def upload_bytes_to_drive(file_name, file_bytes, parent_folder_id, mime_type):
    service = get_drive_service()
    media = MediaIoBaseUpload(io.BytesIO(file_bytes), mimetype=mime_type, resumable=False)
    metadata = {"name": file_name}
    if parent_folder_id:
        metadata["parents"] = [parent_folder_id]
    created = service.files().create(
        body=metadata,
        media_body=media,
        fields="id, webViewLink",
        supportsAllDrives=True,
    ).execute()
    return {
        "file_id": created.get("id"),
        "web_view_link": created.get("webViewLink"),
    }


class Camp(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(200), nullable=False, unique=True, index=True)
    created_at = db.Column(db.DateTime, default=datetime.datetime.utcnow)
    works = db.relationship("Work", backref="camp", lazy=True)
    batches = db.relationship("GenerationBatch", backref="camp", lazy=True)


class GenerationBatch(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    session_name = db.Column(db.String(255), default="", index=True)
    contract_date = db.Column(db.Date, nullable=False)
    created_by = db.Column(db.String(100), default="")
    status = db.Column(db.String(50), default="draft")
    created_at = db.Column(db.DateTime, default=datetime.datetime.utcnow)
    works = db.relationship("Work", backref="batch", lazy=True)
    documents = db.relationship("ContractDocument", backref="batch", lazy=True)


class Writer(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    first_name = db.Column(db.String(100), default="", index=True)
    middle_name = db.Column(db.String(100), default="")
    last_names = db.Column(db.String(150), default="")
    full_name = db.Column(db.String(250), nullable=False, unique=True, index=True)
    writer_aka = db.Column(db.String(250), default="")
    ipi = db.Column(db.String(50), nullable=True, unique=True, index=True)
    pro = db.Column(db.String(20), default="")
    email = db.Column(db.String(255), nullable=True, index=True)
    address = db.Column(db.String(255), default="")
    city = db.Column(db.String(100), default="")
    state = db.Column(db.String(100), default="")
    zip_code = db.Column(db.String(20), default="")
    has_master_contract = db.Column(db.Boolean, default=False)
    created_at = db.Column(db.DateTime, default=datetime.datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.datetime.utcnow, onupdate=datetime.datetime.utcnow)


class Work(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(255), nullable=False, index=True)
    normalized_title = db.Column(db.String(255), index=True, default="")
    camp_id = db.Column(db.Integer, db.ForeignKey("camp.id"), nullable=True)
    batch_id = db.Column(db.Integer, db.ForeignKey("generation_batch.id"), nullable=True)
    contract_date = db.Column(db.Date, nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.datetime.utcnow)
    work_writers = db.relationship("WorkWriter", backref="work", lazy=True, cascade="all, delete-orphan")
    contract_documents = db.relationship("ContractDocument", backref="work", lazy=True, cascade="all, delete-orphan")


class WorkWriter(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    work_id = db.Column(db.Integer, db.ForeignKey("work.id"), nullable=False)
    writer_id = db.Column(db.Integer, db.ForeignKey("writer.id"), nullable=False)
    writer_percentage = db.Column(db.Float, default=0.0)
    publisher = db.Column(db.String(255), default="")
    publisher_ipi = db.Column(db.String(50), default="")
    publisher_address = db.Column(db.String(255), default="")
    publisher_city = db.Column(db.String(100), default="")
    publisher_state = db.Column(db.String(100), default="")
    publisher_zip_code = db.Column(db.String(20), default="")
    created_at = db.Column(db.DateTime, default=datetime.datetime.utcnow)
    writer = db.relationship("Writer", backref="work_links")


class ContractDocument(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    batch_id = db.Column(db.Integer, db.ForeignKey("generation_batch.id"), nullable=True)
    work_id = db.Column(db.Integer, db.ForeignKey("work.id"), nullable=True)
    writer_id = db.Column(db.Integer, db.ForeignKey("writer.id"), nullable=False)
    document_type = db.Column(db.String(50), nullable=False)
    file_name = db.Column(db.String(255), nullable=False)
    writer_name_snapshot = db.Column(db.String(250), nullable=False)
    work_title_snapshot = db.Column(db.String(255), nullable=False)
    drive_file_id = db.Column(db.String(255), nullable=True)
    drive_web_view_link = db.Column(db.String(500), nullable=True)
    signed_file_name = db.Column(db.String(255), nullable=True)
    signed_drive_file_id = db.Column(db.String(255), nullable=True)
    signed_web_view_link = db.Column(db.String(500), nullable=True)
    signed_uploaded_at = db.Column(db.DateTime, nullable=True)
    status = db.Column(db.String(50), default="generated")
    generated_at = db.Column(db.DateTime, default=datetime.datetime.utcnow)
    writer = db.relationship("Writer", backref="contract_documents")
    docusign_envelope_id = db.Column(db.String(100), nullable=True)
    docusign_status = db.Column(db.String(50), nullable=True)
    sent_for_signature_at = db.Column(db.DateTime, nullable=True)
    completed_at = db.Column(db.DateTime, nullable=True)
    signed_pdf_drive_file_id = db.Column(db.String(255), nullable=True)
    signed_pdf_drive_web_view_link = db.Column(db.String(500), nullable=True)
    certificate_drive_file_id = db.Column(db.String(255), nullable=True)
    certificate_drive_web_view_link = db.Column(db.String(500), nullable=True)


def init_db():
    with app.app_context():
        db.create_all()


@app.context_processor
def inject_globals():
    return {"team_auth_enabled": bool(TEAM_USERNAME and TEAM_PASSWORD)}


# ================================================================
# CSS
# ================================================================

_STYLE = """
<link rel="preconnect" href="https://fonts.googleapis.com">
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
<link href="https://fonts.googleapis.com/css2?family=DM+Sans:ital,opsz,wght@0,9..40,300;0,9..40,400;0,9..40,500;0,9..40,600;0,9..40,700;1,9..40,400&family=DM+Mono:wght@400;500&display=swap" rel="stylesheet">
<style>
*,*::before,*::after{box-sizing:border-box;margin:0;padding:0}
:root{
  --bg0:#070b12;--bg1:#0c1120;--bg2:#101724;--bg3:#0a0e19;--bg4:#161f32;--bg5:#1c2740;
  --b0:rgba(255,255,255,.07);--b1:rgba(255,255,255,.04);--bf:rgba(99,133,255,.5);
  --a:#6385ff;--ae:#a55bff;--ag:#34d399;--ar:#ff4f6a;--am:#f59e0b;--ac:#22d3ee;
  --t1:#edf0f8;--t2:#8a96b0;--t3:#4a5470;
  --rs:7px;--rm:11px;--rl:15px;
  --sb:230px;--sb-collapsed:54px;--tb:54px;
  --sh:0 4px 28px rgba(0,0,0,.45);
  --f:'DM Sans',system-ui,sans-serif;
  --fm:'DM Mono','Fira Mono',monospace;
}
html,body{height:100%;background:var(--bg0);color:var(--t1);font-family:var(--f);font-size:15px;line-height:1.55;-webkit-font-smoothing:antialiased}
.app{display:flex;min-height:100vh}
.main{margin-left:var(--sb);flex:1;min-height:100vh;transition:margin-left .22s ease}
.page{max-width:1200px;margin:0 auto;padding:22px 22px 100px}
.sb{width:var(--sb);min-height:100vh;background:var(--bg1);border-right:1px solid var(--b0);display:flex;flex-direction:column;position:fixed;left:0;top:0;z-index:50;transition:width .22s ease;overflow:hidden}
.sb.collapsed{width:var(--sb-collapsed)}
.app.sb-collapsed .main{margin-left:var(--sb-collapsed)}
.sb-logo{display:flex;align-items:center;gap:10px;padding:15px 13px 13px;border-bottom:1px solid var(--b0);margin-bottom:5px;text-decoration:none;white-space:nowrap;overflow:hidden}
.sb-ico{width:28px;height:28px;background:linear-gradient(135deg,var(--a),var(--ae));border-radius:7px;display:flex;align-items:center;justify-content:center;font-size:13px;flex-shrink:0}
.sb-name{font-size:14px;font-weight:700;color:var(--t1);letter-spacing:-.02em;transition:opacity .18s}
.sb.collapsed .sb-name{opacity:0;pointer-events:none}
.sb-toggle{display:flex;align-items:center;justify-content:center;width:28px;height:28px;background:var(--bg4);border:1px solid var(--b0);border-radius:6px;cursor:pointer;color:var(--t3);font-size:11px;margin-left:auto;flex-shrink:0;transition:color .14s,background .14s;user-select:none}
.sb-toggle:hover{color:var(--t1);background:var(--bg5)}
.sb.collapsed .sb-toggle{margin-left:0}
.sb-sec{font-size:9.5px;font-weight:700;letter-spacing:.11em;text-transform:uppercase;color:var(--t3);padding:13px 14px 4px;white-space:nowrap;overflow:hidden;transition:opacity .18s}
.sb.collapsed .sb-sec{opacity:0;height:0;padding:0;pointer-events:none}
.sb-nav a{display:flex;align-items:center;gap:9px;padding:8px 13px;color:var(--t2);text-decoration:none;font-size:13px;font-weight:500;transition:color .14s,background .14s;position:relative;white-space:nowrap;overflow:hidden}
.sb-nav a:hover{color:var(--t1);background:rgba(255,255,255,.03)}
.sb-nav a.on{color:var(--a);background:rgba(99,133,255,.08)}
.sb-nav a.on::before{content:'';position:absolute;left:0;top:6px;bottom:6px;width:2px;background:var(--a);border-radius:0 2px 2px 0}
.sb-nav .ni{font-size:13px;flex-shrink:0;opacity:.85;min-width:18px;text-align:center}
.sb-nav .nl{transition:opacity .18s}
.sb.collapsed .sb-nav .nl{opacity:0}
.sb-foot{margin-top:auto;padding:13px 14px;border-top:1px solid var(--b0);font-size:11px;color:var(--t3);white-space:nowrap;overflow:hidden;transition:opacity .18s}
.sb-foot b{color:var(--t2);font-size:11.5px;display:block;margin-bottom:2px}
.sb.collapsed .sb-foot{opacity:0;pointer-events:none}
.topbar{position:sticky;top:0;z-index:40;background:rgba(7,11,18,.9);backdrop-filter:blur(16px);-webkit-backdrop-filter:blur(16px);border-bottom:1px solid var(--b0);height:var(--tb);display:flex;align-items:center;padding:0 22px;gap:12px}
.tb-search{display:flex;align-items:center;gap:8px;background:var(--bg3);border:1px solid var(--b0);border-radius:var(--rs);padding:6px 11px;width:220px;color:var(--t3);font-size:12px}
.tb-kbd{margin-left:auto;font-size:10px;font-family:var(--fm);opacity:.45}
.tb-right{display:flex;align-items:center;gap:7px;margin-left:auto}
.pill-group{display:flex;gap:3px}
.pill{padding:5px 12px;border-radius:var(--rs);font-size:12.5px;font-weight:500;text-decoration:none;color:var(--t2);border:1px solid transparent;transition:all .14s}
.pill:hover{color:var(--t1);background:var(--bg4)}
.pill.on{color:var(--t1);background:var(--bg4);border-color:var(--b0)}
.tb-ibtn{width:31px;height:31px;display:flex;align-items:center;justify-content:center;background:var(--bg4);border:1px solid var(--b0);border-radius:var(--rs);color:var(--t2);cursor:pointer;text-decoration:none;transition:all .14s;font-size:13px}
.tb-ibtn:hover{border-color:var(--bf);color:var(--t1)}
.avatar{width:29px;height:29px;border-radius:50%;background:linear-gradient(135deg,var(--a),var(--ae));display:flex;align-items:center;justify-content:center;font-size:10.5px;font-weight:700;color:#fff;flex-shrink:0}
.ph{display:flex;align-items:flex-start;justify-content:space-between;gap:14px;margin-bottom:22px}
.ph-left{display:flex;align-items:center;gap:12px}
.ph-icon{width:36px;height:36px;background:linear-gradient(135deg,rgba(99,133,255,.16),rgba(165,91,255,.16));border:1px solid rgba(99,133,255,.2);border-radius:var(--rm);display:flex;align-items:center;justify-content:center;font-size:16px;flex-shrink:0}
.ph-title{font-size:20px;font-weight:700;letter-spacing:-.03em;line-height:1.2}
.ph-sub{font-size:12px;color:var(--t2);margin-top:2px}
.ph-actions{display:flex;gap:7px;align-items:center;flex-shrink:0}
.flash-list{margin-bottom:14px}
.flash-item{padding:10px 14px;border-radius:var(--rs);font-size:12.5px;margin-bottom:6px;background:rgba(245,158,11,.1);border:1px solid rgba(245,158,11,.22);color:var(--am)}
.card{background:var(--bg2);border:1px solid var(--b0);border-radius:var(--rl);margin-bottom:12px;box-shadow:var(--sh);overflow:hidden}
.card-hd{display:flex;align-items:center;gap:9px;padding:13px 17px;border-bottom:1px solid var(--b0)}
.card-ico{width:25px;height:25px;background:rgba(99,133,255,.1);border:1px solid rgba(99,133,255,.14);border-radius:6px;display:flex;align-items:center;justify-content:center;font-size:11px;flex-shrink:0}
.card-title{font-size:13px;font-weight:600}
.card-actions{margin-left:auto;display:flex;gap:6px}
.card-body{padding:17px}
.g{display:grid;gap:12px}
.g2{grid-template-columns:1fr 1fr}
.g3{grid-template-columns:1fr 1fr 1fr}
.g4{grid-template-columns:1fr 1fr 1fr 1fr}
.g5{grid-template-columns:1fr 1.5fr .75fr .75fr 1.5fr}
.g52{grid-template-columns:1fr 2fr 1fr .55fr .55fr}
.g4a{grid-template-columns:2fr 1fr .55fr .55fr}
.field{display:flex;flex-direction:column;gap:5px}
.label{font-size:10px;font-weight:700;letter-spacing:.07em;text-transform:uppercase;color:var(--t2)}
.inp{background:var(--bg3);border:1px solid var(--b0);border-radius:var(--rs);color:var(--t1);font-family:var(--f);font-size:14px;padding:9px 12px;width:100%;outline:none;transition:border-color .14s,box-shadow .14s;-webkit-appearance:none;appearance:none}
.inp::placeholder{color:var(--t3)}
.inp:focus{border-color:var(--bf);box-shadow:0 0 0 3px rgba(99,133,255,.1)}
select.inp{background-image:url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' width='11' height='11' viewBox='0 0 24 24' fill='none' stroke='%234a5470' stroke-width='2.5'%3E%3Cpolyline points='6 9 12 15 18 9'/%3E%3C/svg%3E");background-repeat:no-repeat;background-position:right 9px center;padding-right:28px;cursor:pointer}
select.inp option{background:var(--bg2);color:var(--t1)}
.inp-wrap{position:relative}
.inp-ico{position:absolute;left:9px;top:50%;transform:translateY(-50%);font-size:12px;color:var(--t3);pointer-events:none}
.inp-wrap .inp{padding-left:28px}
.btn{display:inline-flex;align-items:center;gap:6px;padding:9px 16px;border-radius:var(--rs);font-family:var(--f);font-size:13.5px;font-weight:600;cursor:pointer;border:1px solid transparent;text-decoration:none;transition:all .15s;white-space:nowrap;line-height:1}
.btn-primary{background:linear-gradient(135deg,var(--a),var(--ae));color:#fff;border:none;box-shadow:0 2px 14px rgba(99,133,255,.28)}
.btn-primary:hover{transform:translateY(-1px);box-shadow:0 5px 22px rgba(99,133,255,.42)}
.btn-primary:active{transform:translateY(0)}
.btn-sec{background:transparent;color:var(--t2);border-color:var(--b0)}
.btn-sec:hover{color:var(--t1);border-color:rgba(255,255,255,.14);background:var(--bg4)}
.btn-danger{background:rgba(255,79,106,.1);color:var(--ar);border-color:rgba(255,79,106,.2)}
.btn-danger:hover{background:rgba(255,79,106,.18);border-color:rgba(255,79,106,.38)}
.btn-success{background:rgba(52,211,153,.1);color:var(--ag);border-color:rgba(52,211,153,.22)}
.btn-success:hover{background:rgba(52,211,153,.18);border-color:rgba(52,211,153,.4)}
.btn-cyan{background:rgba(34,211,238,.1);color:var(--ac);border-color:rgba(34,211,238,.22)}
.btn-cyan:hover{background:rgba(34,211,238,.18);border-color:rgba(34,211,238,.4)}
.btn-sm{padding:5px 11px;font-size:12px}
.btn-xs{padding:3px 8px;font-size:11px}
.tag{display:inline-flex;align-items:center;padding:2px 7px;border-radius:99px;font-size:10px;font-weight:700;white-space:nowrap}
.tag-new{background:rgba(245,158,11,.12);color:var(--am);border:1px solid rgba(245,158,11,.2)}
.tag-exist{background:rgba(52,211,153,.12);color:var(--ag);border:1px solid rgba(52,211,153,.2)}
.tag-full{background:rgba(99,133,255,.12);color:var(--a);border:1px solid rgba(99,133,255,.2)}
.tag-s1{background:rgba(52,211,153,.12);color:var(--ag);border:1px solid rgba(52,211,153,.2)}
.status{display:inline-flex;align-items:center;gap:4px;padding:3px 8px;border-radius:99px;font-size:10.5px;font-weight:700}
.status-dot{width:5px;height:5px;border-radius:50%;background:currentColor;flex-shrink:0}
.s-draft{background:rgba(255,255,255,.06);color:var(--t2);border:1px solid var(--b0)}
.s-generated{background:rgba(99,133,255,.12);color:var(--a);border:1px solid rgba(99,133,255,.2)}
.s-sent{background:rgba(34,211,238,.12);color:var(--ac);border:1px solid rgba(34,211,238,.2)}
.s-delivered{background:rgba(245,158,11,.12);color:var(--am);border:1px solid rgba(245,158,11,.2)}
.s-completed,.s-signed,.s-signed_uploaded,.s-signed_complete{background:rgba(52,211,153,.12);color:var(--ag);border:1px solid rgba(52,211,153,.2)}
.s-signed_partial{background:rgba(245,158,11,.12);color:var(--am);border:1px solid rgba(245,158,11,.2)}
.tbl-wrap{overflow-x:auto}
.tbl{width:100%;border-collapse:collapse}
.tbl th{font-size:10px;font-weight:700;letter-spacing:.08em;text-transform:uppercase;color:var(--t3);padding:9px 13px;text-align:left;border-bottom:1px solid var(--b0);white-space:nowrap}
.tbl td{padding:11px 13px;font-size:13px;color:var(--t1);border-bottom:1px solid var(--b1);vertical-align:middle}
.tbl tr:last-child td{border-bottom:none}
.tbl tbody tr:hover td{background:rgba(255,255,255,.02)}
.tbl .empty td{color:var(--t3);text-align:center;padding:26px;font-size:13px}
.tbl a{color:var(--a);text-decoration:none}
.tbl a:hover{text-decoration:underline}
.split-banner{background:var(--bg2);border:1px solid var(--b0);border-radius:var(--rl);padding:13px 17px;margin-bottom:14px;display:flex;align-items:center;gap:16px;position:sticky;top:var(--tb);z-index:30;backdrop-filter:blur(12px)}
.split-info{flex:1;min-width:0}
.split-lr{display:flex;justify-content:space-between;align-items:center;margin-bottom:6px}
.split-lbl{font-size:10px;font-weight:700;letter-spacing:.07em;text-transform:uppercase;color:var(--t2)}
.split-pct{font-family:var(--fm);font-size:12px;font-weight:600;color:var(--t1)}
.split-track{height:5px;background:rgba(255,255,255,.05);border-radius:99px;overflow:hidden}
.split-fill{height:100%;border-radius:99px;transition:width .35s ease,background .35s ease;width:0;background:linear-gradient(90deg,var(--a),var(--ae))}
.split-badge{display:inline-flex;align-items:center;gap:5px;padding:5px 11px;border-radius:99px;font-size:11px;font-weight:700;white-space:nowrap;transition:all .3s}
.sb-inc{background:rgba(245,158,11,.1);color:var(--am);border:1px solid rgba(245,158,11,.2)}
.sb-ok{background:rgba(52,211,153,.1);color:var(--ag);border:1px solid rgba(52,211,153,.2)}
.sb-over{background:rgba(255,79,106,.1);color:var(--ar);border:1px solid rgba(255,79,106,.2)}
.sb-dot{width:5px;height:5px;border-radius:50%;background:currentColor}
.wc{background:var(--bg4);border:1px solid var(--b0);border-radius:var(--rm);margin-bottom:9px;overflow:hidden;transition:border-color .2s}
.wc:hover{border-color:rgba(99,133,255,.2)}
.wc-hd{display:flex;align-items:center;gap:10px;padding:11px 14px;cursor:pointer;border-bottom:1px solid transparent;transition:background .14s,border-color .2s;user-select:none}
.wc-hd:hover{background:rgba(255,255,255,.02)}
.wc.open .wc-hd{border-bottom-color:var(--b0)}
.wc-av{width:29px;height:29px;border-radius:50%;background:linear-gradient(135deg,rgba(99,133,255,.22),rgba(165,91,255,.22));border:1px solid rgba(99,133,255,.18);display:flex;align-items:center;justify-content:center;font-size:12px;color:var(--a);flex-shrink:0}
.wc-nw{flex:1;min-width:0}
.wc-name{font-size:13px;font-weight:600;white-space:nowrap;overflow:hidden;text-overflow:ellipsis}
.wc-sub{font-size:11px;color:var(--t3);margin-top:1px}
.wc-tags{display:flex;align-items:center;gap:5px}
.wc-chev{font-size:9px;color:var(--t3);transition:transform .2s}
.wc.open .wc-chev{transform:rotate(180deg)}
.wc-body{display:none;padding:14px}
.wc.open .wc-body{display:block}
.wc-sec{font-size:9.5px;font-weight:700;letter-spacing:.09em;text-transform:uppercase;color:var(--t3);padding-bottom:7px;margin-bottom:10px;border-bottom:1px solid var(--b0);margin-top:13px}
.wc-sec:first-child{margin-top:0}
.ac-wrap{position:relative}
.ac-box{position:absolute;top:calc(100% + 3px);left:0;right:0;z-index:200;background:var(--bg2);border:1px solid var(--b0);border-radius:var(--rs);max-height:190px;overflow-y:auto;display:none;box-shadow:0 8px 36px rgba(0,0,0,.55)}
.ac-item{padding:8px 12px;cursor:pointer;border-bottom:1px solid var(--b1);transition:background .1s}
.ac-item:last-child{border-bottom:none}
.ac-item:hover{background:var(--bg4)}
.ac-item strong{color:var(--t1);font-size:12.5px}
.ac-item small{color:var(--t3);font-size:11px}
.action-bar{position:fixed;bottom:0;left:var(--sb);right:0;background:rgba(7,11,18,.94);backdrop-filter:blur(18px);-webkit-backdrop-filter:blur(18px);border-top:1px solid var(--b0);padding:12px 22px;display:flex;align-items:center;gap:9px;z-index:45;transition:left .22s ease}
.app.sb-collapsed .action-bar{left:var(--sb-collapsed)}
.ab-space{flex:1}
.upl-form{display:flex;gap:6px;align-items:center}
.upl-inp{background:var(--bg3);border:1px solid var(--b0);border-radius:var(--rs);color:var(--t2);font-size:11px;font-family:var(--f);padding:4px 7px;cursor:pointer;flex:1;min-width:0;max-width:160px}
.upl-inp::-webkit-file-upload-button{background:var(--bg4);border:1px solid var(--b0);border-radius:5px;color:var(--t2);font-family:var(--f);font-size:10.5px;padding:3px 7px;cursor:pointer;margin-right:6px}
.spin{display:none;width:11px;height:11px;border:2px solid rgba(255,255,255,.25);border-top-color:#fff;border-radius:50%;animation:spin .6s linear infinite}
.spin.on{display:inline-block}
@keyframes spin{to{transform:rotate(360deg)}}
.info-grid{display:grid;grid-template-columns:1fr 1fr;gap:8px 22px}
.info-item label{font-size:10px;font-weight:700;letter-spacing:.06em;text-transform:uppercase;color:var(--t3);display:block;margin-bottom:1px}
.info-item span,.info-item a{font-size:13px;color:var(--t1)}
.info-item a{color:var(--a);text-decoration:none}
.info-item a:hover{text-decoration:underline}
.file-link{color:var(--a);text-decoration:none;font-size:11.5px;display:inline-flex;align-items:center;gap:4px;max-width:200px;overflow:hidden;text-overflow:ellipsis;white-space:nowrap}
.file-link:hover{text-decoration:underline}
.file-link-plain{color:var(--t2);font-size:11.5px;max-width:200px;overflow:hidden;text-overflow:ellipsis;white-space:nowrap;display:inline-block}
::-webkit-scrollbar{width:5px;height:5px}
::-webkit-scrollbar-track{background:transparent}
::-webkit-scrollbar-thumb{background:rgba(255,255,255,.07);border-radius:99px}
.login-wrap{min-height:100vh;display:flex;align-items:center;justify-content:center;background:var(--bg0);padding:20px}
.login-card{width:100%;max-width:360px;background:var(--bg2);border:1px solid var(--b0);border-radius:var(--rl);padding:36px 32px;box-shadow:var(--sh)}
.login-logo{display:flex;align-items:center;gap:10px;margin-bottom:28px}
.login-logo-ico{width:34px;height:34px;background:linear-gradient(135deg,var(--a),var(--ae));border-radius:9px;display:flex;align-items:center;justify-content:center;font-size:16px}
.login-logo-name{font-size:17px;font-weight:700;letter-spacing:-.02em}
.login-h{font-size:17px;font-weight:700;margin-bottom:4px}
.login-sub{font-size:12.5px;color:var(--t2);margin-bottom:22px}
.login-field{margin-bottom:14px}
@media(max-width:860px){.g3{grid-template-columns:1fr 1fr}.g5,.g52{grid-template-columns:1fr 1fr}.g4,.g4a{grid-template-columns:1fr 1fr}}
@media(max-width:640px){.sb{display:none}.main{margin-left:0!important}.page{padding:16px 13px 90px}.g3,.g2,.g4,.g4a,.g5,.g52{grid-template-columns:1fr}.topbar{padding:0 13px}.tb-search{display:none}.action-bar{left:0!important;padding:11px 14px}}
/* ===== Dynamic sidebar overrides ===== */
.sb,
.main,
.action-bar{transition:width .22s ease,margin-left .22s ease,left .22s ease}

.sb-toggle{
  display:flex;align-items:center;justify-content:center;
  width:28px;height:28px;
  background:var(--bg4);
  border:1px solid var(--b0);
  border-radius:6px;
  cursor:pointer;
  color:var(--t3);
  font-size:11px;
  margin-left:auto;
  flex-shrink:0;
  transition:color .14s,background .14s,transform .14s;
  user-select:none;
}
.sb-toggle:hover{color:var(--t1);background:var(--bg5)}

.sb.collapsed{
  width:var(--sb-collapsed);
}
.app.sb-collapsed .main{
  margin-left:var(--sb-collapsed);
}
.app.sb-collapsed .action-bar{
  left:var(--sb-collapsed);
}

.sb.collapsed .sb-logo{
  justify-content:center;
  gap:0;
  padding:15px 6px 13px;
}
.sb.collapsed .sb-name{
  opacity:0;
  width:0;
  pointer-events:none;
}
.sb.collapsed .sb-sec{
  opacity:0;
  height:0;
  padding:0;
  pointer-events:none;
}
.sb.collapsed .sb-foot{
  opacity:0;
  pointer-events:none;
}
.sb.collapsed .sb-nav a{
  justify-content:center;
  gap:0;
  padding:10px 0;
}
.sb.collapsed .sb-nav .nl{
  opacity:0;
  width:0;
  pointer-events:none;
}
.sb.collapsed .sb-toggle{
  position:absolute;
  right:8px;
  top:14px;
  margin-left:0;
}

.sb.collapsed.hover-open{
  width:var(--sb);
  box-shadow:16px 0 34px rgba(0,0,0,.34);
}
.sb.collapsed.hover-open .sb-logo{
  justify-content:flex-start;
  gap:10px;
  padding:15px 13px 13px;
}
.sb.collapsed.hover-open .sb-name{
  opacity:1;
  width:auto;
  pointer-events:auto;
}
.sb.collapsed.hover-open .sb-sec{
  opacity:1;
  height:auto;
  padding:13px 14px 4px;
  pointer-events:auto;
}
.sb.collapsed.hover-open .sb-foot{
  opacity:1;
  pointer-events:auto;
}
.sb.collapsed.hover-open .sb-nav a{
  justify-content:flex-start;
  gap:9px;
  padding:8px 13px;
}
.sb.collapsed.hover-open .sb-nav .nl{
  opacity:1;
  width:auto;
  pointer-events:auto;
}

.sb-nav .ni{
  font-size:14px;
  min-width:18px;
  text-align:center;
  flex-shrink:0;
}
.sb-nav .ni-pencil{
  color:var(--am);
}
/* Custom pencil icon (matches your header icon) */
.ni-pencil-custom{
  width:18px;
  height:18px;
  display:inline-flex;
  align-items:center;
  justify-content:center;
  border-radius:6px;
  background:linear-gradient(135deg, rgba(99,133,255,.25), rgba(165,91,255,.25));
  border:1px solid rgba(99,133,255,.25);
  font-size:11px;
  position:relative;
}

/* Pencil emoji inside */
.ni-pencil-custom::before{
  content:"✏️";
  font-size:11px;
  filter:saturate(1.2);
}
</style>"""

# ================================================================
# SHARED SIDEBAR JS
# ================================================================

_SB_JS = """
<script>
(function(){
  function getEls(){
    return {
      sb: document.getElementById('mainSidebar'),
      app: document.getElementById('mainApp'),
      tog: document.getElementById('sbToggle')
    };
  }

  function applySidebarMode(mode){
    var els = getEls();
    if(!els.sb || !els.app) return;

    var collapsed = mode === 'closed';
    els.sb.classList.toggle('collapsed', collapsed);
    els.app.classList.toggle('sb-collapsed', collapsed);

    if(!collapsed){
      els.sb.classList.remove('hover-open');
    }

    if(els.tog){
      els.tog.textContent = collapsed ? '>' : '<';
      els.tog.title = collapsed ? 'Pin sidebar open' : 'Pin sidebar closed';
    }
  }

  window.toggleSidebar = function(e){
    if(e){
      e.preventDefault();
      e.stopPropagation();
    }
    var current = localStorage.getItem('sb_mode') || 'open';
    var next = current === 'closed' ? 'open' : 'closed';
    localStorage.setItem('sb_mode', next);
    applySidebarMode(next);
  };

  document.addEventListener('DOMContentLoaded', function(){
    var els = getEls();
    if(!els.sb || !els.app) return;

    var savedMode = localStorage.getItem('sb_mode') || 'open';
    applySidebarMode(savedMode);

    els.sb.addEventListener('mouseenter', function(){
      var mode = localStorage.getItem('sb_mode') || 'open';
      if(mode === 'closed'){
        els.sb.classList.add('hover-open');
      }
    });

    els.sb.addEventListener('mouseleave', function(){
      var mode = localStorage.getItem('sb_mode') || 'open';
      if(mode === 'closed'){
        els.sb.classList.remove('hover-open');
      }
    });
  });
})();
</script>"""

# ================================================================
# SHARED SIDEBAR HTML
# ================================================================

def _sidebar(active):
    pages = [
        ("works_list",   "Works",     "<span class='ni'>&#127925;</span>"),
        ("formulario",   "New Work",  "<span class='ni ni-pencil-custom'></span>"),
        ("batches_list", "Sessions",  "<span class='ni'>&#128230;</span>"),
    ]

    html = "<aside class='sb' id='mainSidebar'>"
    html += "<a class='sb-logo' href='/works'>"
    html += "<div class='sb-ico'>&#127925;</div>"
    html += "<span class='sb-name'>LabelMind</span>"
    html += "<span class='sb-toggle' id='sbToggle' onclick='toggleSidebar(event)' title='Pin sidebar closed'>&lt;</span>"
    html += "</a>"

    html += "<div class='sb-sec'>Contracts</div>"
    html += "<nav class='sb-nav'>"

    for endpoint, label, icon_html in pages:
        on = " class='on'" if active == endpoint else ""
        if endpoint == "works_list":
            href = "/works"
        elif endpoint == "formulario":
            href = "/"
        else:
            href = "/batches"

        html += "<a href='" + href + "'" + on + " title='" + label + "'>"
        html += icon_html
        html += "<span class='nl'>" + label + "</span></a>"

    html += "<a href='#' title='Templates'><span class='ni'>&#128196;</span><span class='nl'>Templates</span></a>"
    html += "</nav>"

    html += "<div class='sb-sec'>Resources</div>"
    html += "<nav class='sb-nav'>"
    html += "<a href='#' title='Writer Directory'><span class='ni'>&#128101;</span><span class='nl'>Writer Directory</span></a>"
    html += "<a href='#' title='Settings'><span class='ni'>&#9881;</span><span class='nl'>Settings</span></a>"
    html += "</nav>"

    html += "<div class='sb-foot'><b>LabelMind</b>Music Publishing Contracts<br>2026 LabelMind.ai</div>"
    html += "</aside>"
    return html

def _topbar(pill=""):
    works_on = " class='pill on'" if pill == "works" else " class='pill'"
    sess_on = " class='pill on'" if pill == "sessions" else " class='pill'"
    html = "<header class='topbar'>"
    html += "<div class='tb-search'>Search works, sessions, writers...</div>"
    html += "<div class='tb-right'>"
    html += "<div class='pill-group'>"
    html += "<a href='/works'" + works_on + ">Works</a>"
    html += "<a href='/batches'" + sess_on + ">Sessions</a>"
    html += "</div>"
    html += "{% if team_auth_enabled and session.get('logged_in') %}"
    html += "<a href='/logout' class='tb-ibtn' title='Log out'>&#128682;</a>"
    html += "{% endif %}"
    html += "<div class='avatar'>IS</div>"
    html += "</div></header>"
    return html


# ================================================================
# LOGIN
# ================================================================

LOGIN_HTML = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>Login - LabelMind</title>""" + _STYLE + """
</head>
<body>
<div class="login-wrap"><div class="login-card">
<div class="login-logo">
  <div class="login-logo-ico">&#127925;</div>
  <span class="login-logo-name">LabelMind</span>
</div>
<div class="login-h">Welcome back</div>
<div class="login-sub">Music Publishing Contracts</div>
{% with messages = get_flashed_messages() %}
{% if messages %}
<div class="flash-list">{% for m in messages %}<div class="flash-item">&#9888; {{ m }}</div>{% endfor %}</div>
{% endif %}{% endwith %}
<form method="post">
<div class="login-field"><label class="label">Username</label>
<input class="inp" name="username" required autocomplete="username" placeholder="your username"></div>
<div class="login-field"><label class="label">Password</label>
<input class="inp" type="password" name="password" required autocomplete="current-password" placeholder="password"></div>
<button class="btn btn-primary" style="width:100%;justify-content:center;margin-top:4px;">Log in</button>
</form>
</div></div>
</body></html>"""

# ================================================================
# NEW WORK FORM
# ================================================================

FORM_HTML = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>New Work - LabelMind</title>""" + _STYLE + """
</head>
<body>
<div class="app" id="mainApp">
""" + _sidebar("formulario") + """
<main class="main">
""" + _topbar("") + """
<div class="page">
{% with messages = get_flashed_messages() %}
{% if messages %}
<div class="flash-list">{% for m in messages %}<div class="flash-item">&#9888; {{ m }}</div>{% endfor %}</div>
{% endif %}{% endwith %}
<div class="ph">
  <div class="ph-left">
    <div class="ph-icon">&#9999;</div>
    <div>
      <div class="ph-title">New Work</div>
      <div class="ph-sub">Create a work, add writers, and save into a session for contract generation.</div>
    </div>
  </div>
</div>
<div class="split-banner">
  <div class="split-info">
    <div class="split-lr">
      <span class="split-lbl">Split Total</span>
      <span class="split-pct"><span id="splitTotal">0.00</span>% Complete</span>
    </div>
    <div class="split-track"><div class="split-fill" id="splitFill"></div></div>
  </div>
  <div class="split-badge sb-inc" id="splitBadge"><span class="sb-dot"></span>Incomplete</div>
</div>
<form method="post" id="workForm">
  <input type="hidden" name="force_create" value="{{ force_create or '' }}">
  <div class="card">
    <div class="card-hd"><div class="card-ico">&#128203;</div><span class="card-title">Work Information</span></div>
    <div class="card-body">
      <div class="g g2" style="margin-bottom:12px">
  <div class="field">
    <label class="label">Add to Existing Session</label>
    <div class="inp-wrap">
      <span class="inp-ico">&#128230;</span>
      <select class="inp" name="existing_batch_id">
        <option value="">-- Create new session</option>
        {% for batch in batches %}
        <option value="{{ batch.id }}" {% if selected_batch_id == (batch.id|string) %}selected{% endif %}>
          Session #{{ batch.id }}{% if batch.session_name %} -- {{ batch.session_name }}{% endif %} -- {{ batch.contract_date.strftime('%Y-%m-%d') }}
        </option>
        {% endfor %}
      </select>
    </div>
  </div>
  <div class="field">
    <label class="label">Create New Session</label>
    <input class="inp" name="new_session_name" placeholder="Session Name">
  </div>
</div>
      <div class="g g2">
        <div class="field">
          <label class="label">Work Title</label>
          <div class="inp-wrap">
            <span class="inp-ico">&#127925;</span>
            <input class="inp" name="work_title" required placeholder="e.g. La Serenata">
          </div>
        </div>
        <div class="field">
          <label class="label">Contract Date</label>
          <div class="inp-wrap">
            <span class="inp-ico">&#128197;</span>
            <input class="inp" name="contract_date" type="date" required>
          </div>
        </div>
      </div>
    </div>
  </div>
  <div class="card">
    <div class="card-hd">
      <div class="card-ico">&#128101;</div>
      <span class="card-title">Writers</span>
      <div class="card-actions">
        <button type="button" class="btn btn-sec btn-sm" onclick="addWriter()">+ Add Writer</button>
      </div>
    </div>
    <div class="card-body" id="writerRows" style="padding-bottom:6px"></div>
  </div>
  <div class="action-bar">
    <button type="button" class="btn btn-sec" onclick="addWriter()">+ Add Writer</button>
    <div class="ab-space"></div>
    <button type="submit" class="btn btn-primary">Save Work to Session</button>
  </div>
</form>
</div>
</main>
</div>
""" + _SB_JS + """
<script>
var proMap = {
  BMI: {name:'Songs of Afinarte', ipi:'817874992'},
  ASCAP: {name:'Melodies of Afinarte', ipi:'807953316'},
  SESAC: {name:'Music of Afinarte', ipi:'817094629'}
};
var defAddr = "{{ default_publisher_address }}";
var defCity = "{{ default_publisher_city }}";
var defState = "{{ default_publisher_state }}";
var defZip = "{{ default_publisher_zip }}";
var idx = 0;

function statusHtml(ws, ct) {
  var wc = ws === 'Existing Writer' ? 'tag-exist' : 'tag-new';
  var cc = ct === 'Schedule 1' ? 'tag-s1' : 'tag-full';
  return '<span class="tag ' + wc + '">' + ws + '</span><span class="tag ' + cc + '">' + ct + '</span>';
}

function writerTpl(i) {
  var h = '<div class="wc open" data-idx="' + i + '">';
  h += '<div class="wc-hd" onclick="toggleWC(this)">';
  h += '<div class="wc-av">&#128100;</div>';
  h += '<div class="wc-nw"><div class="wc-name wc-dn">Writer ' + (i + 1) + '</div>';
  h += '<div class="wc-sub wc-ds">New</div></div>';
  h += '<div class="wc-tags wc-meta">' + statusHtml('New Writer', 'Full Contract') + '</div>';
  h += '<span class="wc-chev">v</span>';
  h += '<button type="button" class="btn btn-danger btn-sm" onclick="rmWriter(event,this)" style="margin-left:8px">Remove</button>';
  h += '</div>';
  h += '<div class="wc-body">';
  h += '<input type="hidden" name="writer_id" class="wid">';
  h += '<div class="wc-sec">Identity</div>';
  h += '<div class="g g4" style="gap:10px">';
  h += '<div class="field ac-wrap"><label class="label">First Name</label>';
  h += '<input class="inp wfn" name="writer_first_name" placeholder="First" autocomplete="off">';
  h += '<div class="ac-box wsug"></div></div>';
  h += '<div class="field"><label class="label">Middle Name</label>';
  h += '<input class="inp wmn" name="writer_middle_name" placeholder="Middle" autocomplete="off"></div>';
  h += '<div class="field"><label class="label">Last Name(s)</label>';
  h += '<input class="inp wln" name="writer_last_names" placeholder="Last Name" autocomplete="off"></div>';
  h += '<div class="field"><label class="label">AKA / Stage</label>';
  h += '<input class="inp waka" name="writer_aka" placeholder="Stage Name"></div>';
  h += '</div>';
  h += '<div class="wc-sec">Publishing</div>';
  h += '<div class="g g5" style="gap:10px">';
  h += '<div class="field"><label class="label">IPI #</label>';
  h += '<input class="inp wipi" name="writer_ipi" placeholder="IPI Number"></div>';
  h += '<div class="field"><label class="label">Email</label>';
  h += '<input class="inp wem" name="writer_email" placeholder="writer@email.com" type="email"></div>';
  h += '<div class="field"><label class="label">PRO</label>';
  h += '<select class="inp wpro" name="writer_pro" onchange="syncPro(this)">';
  h += '<option value="">PRO</option><option value="BMI">BMI</option>';
  h += '<option value="ASCAP">ASCAP</option><option value="SESAC">SESAC</option></select></div>';
  h += '<div class="field"><label class="label">Writer %</label>';
  h += '<input class="inp wspl" name="writer_percentage" placeholder="0" type="number" step="0.01" min="0" max="100"></div>';
  h += '<div class="field"><label class="label">Publisher</label>';
  h += '<input class="inp wpub" name="writer_publisher" placeholder="Publisher Name"></div>';
  h += '</div>';
  h += '<div class="wc-sec">Publisher Details</div>';
  h += '<div class="g g52" style="gap:10px">';
  h += '<div class="field"><label class="label">Publisher IPI</label>';
  h += '<input class="inp wpipi" name="publisher_ipi" placeholder="Publisher IPI"></div>';
  h += '<div class="field"><label class="label">Address</label>';
  h += '<input class="inp wpaddr" name="publisher_address" value="' + defAddr + '" placeholder="Address"></div>';
  h += '<div class="field"><label class="label">City</label>';
  h += '<input class="inp wpcity" name="publisher_city" value="' + defCity + '" placeholder="City"></div>';
  h += '<div class="field"><label class="label">State</label>';
  h += '<input class="inp wpst" name="publisher_state" value="' + defState + '" placeholder="ST"></div>';
  h += '<div class="field"><label class="label">Zip</label>';
  h += '<input class="inp wpzip" name="publisher_zip_code" value="' + defZip + '" placeholder="Zip"></div>';
  h += '</div>';
  h += '<div class="wc-sec">Writer Address</div>';
  h += '<div class="g g4a" style="gap:10px">';
  h += '<div class="field"><label class="label">Street</label>';
  h += '<input class="inp waddr" name="writer_address" placeholder="Street Address"></div>';
  h += '<div class="field"><label class="label">City</label>';
  h += '<input class="inp wcity" name="writer_city" placeholder="City"></div>';
  h += '<div class="field"><label class="label">State</label>';
  h += '<input class="inp wst" name="writer_state" placeholder="ST"></div>';
  h += '<div class="field"><label class="label">Zip</label>';
  h += '<input class="inp wzip" name="writer_zip_code" placeholder="Zip"></div>';
  h += '</div></div></div>';
  return h;
}

function toggleWC(hd) { hd.closest('.wc').classList.toggle('open'); }

function reindexWriters() {
  document.querySelectorAll('#writerRows .wc').forEach(function(card, i) {
    card.dataset.idx = i;
    var fn = (card.querySelector('.wfn').value || '').trim();
    var mn = (card.querySelector('.wmn').value || '').trim();
    var ln = (card.querySelector('.wln').value || '').trim();
    if (!fn && !mn && !ln) {
      card.querySelector('.wc-dn').textContent = 'Writer ' + (i + 1);
    }
  });
}

function addWriter() {
  var c = document.getElementById('writerRows');
  c.insertAdjacentHTML('beforeend', writerTpl(idx));
  var card = c.lastElementChild;
  setupWriter(card);
  idx++;
  reindexWriters();
  recalc();
}

function rmWriter(e, btn) {
  e.stopPropagation();
  btn.closest('.wc').remove();
  reindexWriters();
  recalc();
}

function syncPro(sel) {
  var r = sel.closest('.wc');
  var p = proMap[sel.value];
  if (!p) return;
  r.querySelector('.wpub').value = p.name;
  r.querySelector('.wpipi').value = p.ipi;
  updateHdr(r);
}

function fullName(r) {
  return [
    r.querySelector('.wfn').value,
    r.querySelector('.wmn').value,
    r.querySelector('.wln').value
  ].map(function(s) { return s.trim(); }).filter(Boolean).join(' ');
}

function updateHdr(r) {
  var i = parseInt(r.dataset.idx) || 0;
  var n = fullName(r) || 'Writer ' + (i + 1);
  var pro = r.querySelector('.wpro').value || '--';
  var pct = r.querySelector('.wspl').value || '--';
  r.querySelector('.wc-dn').textContent = n;
  r.querySelector('.wc-ds').textContent = pro + ' / ' + pct + '%';
}

function setStatus(r, ws, ct) { r.querySelector('.wc-meta').innerHTML = statusHtml(ws, ct); }
function hideSug(r) { var b = r.querySelector('.wsug'); b.style.display = 'none'; b.innerHTML = ''; }
function resetNew(r) { r.querySelector('.wid').value = ''; setStatus(r, 'New Writer', 'Full Contract'); }

function fillWriter(r, w) {
  r.querySelector('.wid').value = w.id || '';
  r.querySelector('.wfn').value = w.first_name || '';
  r.querySelector('.wmn').value = w.middle_name || '';
  r.querySelector('.wln').value = w.last_names || '';
  r.querySelector('.waka').value = w.writer_aka || '';
  r.querySelector('.wipi').value = w.ipi || '';
  r.querySelector('.wem').value = w.email || '';
  r.querySelector('.wpro').value = w.pro || '';
  r.querySelector('.waddr').value = w.address || '';
  r.querySelector('.wcity').value = w.city || '';
  r.querySelector('.wst').value = w.state || '';
  r.querySelector('.wzip').value = w.zip_code || '';
  var pd = proMap[w.pro] || {};
  r.querySelector('.wpub').value = w.default_publisher || pd.name || '';
  r.querySelector('.wpipi').value = w.default_publisher_ipi || pd.ipi || '';
  updateHdr(r);
  setStatus(r, 'Existing Writer', w.has_master_contract ? 'Schedule 1' : 'Full Contract');
  hideSug(r);
}

function setupWriter(r) {
  var fn = r.querySelector('.wfn');
  var mn = r.querySelector('.wmn');
  var ln = r.querySelector('.wln');
  var sug = r.querySelector('.wsug');
  var spl = r.querySelector('.wspl');
  var pro = r.querySelector('.wpro');

  function search() {
    var q = fullName(r);
    if (q.length < 2) { hideSug(r); resetNew(r); return; }
    fetch('/writers/search?q=' + encodeURIComponent(q))
      .then(function(res) { return res.json(); })
      .then(function(ws) {
        if (!ws.length) { hideSug(r); resetNew(r); return; }
        sug.innerHTML = ws.map(function(w) {
          var safeW = JSON.stringify(w).replace(/'/g, "&#39;");
          return "<div class='ac-item' data-w='" + safeW + "'>" +
            "<strong>" + w.full_name + "</strong><br>" +
            "<small>" + (w.city || '') + (w.city && w.state ? ', ' : '') + (w.state || '') + "</small>" +
            "</div>";
        }).join('');
        sug.style.display = 'block';
        sug.querySelectorAll('.ac-item').forEach(function(item) {
          item.addEventListener('click', function() {
            fillWriter(r, JSON.parse(item.dataset.w));
          });
        });
      });
  }

  [fn, mn, ln].forEach(function(inp) {
    inp.addEventListener('input', function() {
      resetNew(r); updateHdr(r); reindexWriters(); search();
    });
  });
  spl.addEventListener('input', function() { updateHdr(r); recalc(); });
  pro.addEventListener('change', function() { updateHdr(r); });
  document.addEventListener('click', function(e) {
    if (![fn, mn, ln, sug].some(function(el) { return el.contains(e.target); })) {
      hideSug(r);
    }
  });
}

function recalc() {
  var total = 0;
  document.querySelectorAll('.wspl').forEach(function(i) {
    total += parseFloat(i.value || 0) || 0;
  });
  var rounded = total.toFixed(2);
  document.getElementById('splitTotal').textContent = rounded;
  var fill = document.getElementById('splitFill');
  var badge = document.getElementById('splitBadge');
  fill.style.width = Math.min(total, 100) + '%';
  if (Math.abs(total - 100) < 0.001) {
    fill.style.background = 'linear-gradient(90deg,#34d399,#059669)';
    badge.className = 'split-badge sb-ok';
    badge.innerHTML = '<span class="sb-dot"></span>Complete';
  } else if (total > 100) {
    fill.style.background = 'linear-gradient(90deg,#ff4f6a,#c0152d)';
    badge.className = 'split-badge sb-over';
    badge.innerHTML = '<span class="sb-dot"></span>Over 100%';
  } else {
    fill.style.background = 'linear-gradient(90deg,var(--a),var(--ae))';
    badge.className = 'split-badge sb-inc';
    badge.innerHTML = '<span class="sb-dot"></span>Incomplete';
  }
}

document.getElementById('workForm').addEventListener('submit', function(e) {
  var rows = document.querySelectorAll('.wc');
  if (!rows.length) { e.preventDefault(); alert('Add at least one writer.'); return; }
  var ok = false;
  for (var i = 0; i < rows.length; i++) {
    var r = rows[i];
    var n = fullName(r);
    var s = parseFloat(r.querySelector('.wspl').value || 0) || 0;
    if (n) { ok = true; if (s <= 0) { e.preventDefault(); alert('Each writer must have a split > 0.'); return; } }
  }
  if (!ok) { e.preventDefault(); alert('Add at least one writer with a name.'); return; }
  var t = parseFloat(document.getElementById('splitTotal').textContent || 0) || 0;
  if (Math.abs(t - 100) >= 0.001) { e.preventDefault(); alert('Total split must equal 100%.'); }
});

addWriter();
</script>
</body></html>"""

# ================================================================
# DUPLICATE WARNING
# ================================================================

DUPLICATE_WARNING_HTML = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>Possible Duplicate - LabelMind</title>""" + _STYLE + """
</head>
<body>
<div class="app" id="mainApp">
""" + _sidebar("formulario") + """
<main class="main">
""" + _topbar("") + """
<div class="page">
<div class="ph">
  <div class="ph-left">
    <div class="ph-icon">&#9888;</div>
    <div>
      <div class="ph-title">Possible Duplicate Found</div>
      <div class="ph-sub">Existing works match this title and writer set.</div>
    </div>
  </div>
</div>
<div class="card">
  <div class="card-hd"><div class="card-ico">&#128269;</div><span class="card-title">Matching Works</span></div>
  <div class="card-body">
    <table class="tbl" style="margin-bottom:18px">
      <thead><tr><th>Title</th><th>Session</th><th>Created</th></tr></thead>
      <tbody>
        {% for item in duplicates %}
        <tr>
          <td style="font-weight:600">{{ item.title }}</td>
          <td>{{ item.camp_name or '--' }}</td>
          <td style="color:var(--t2)">{{ item.created_at }}</td>
        </tr>
        {% endfor %}
      </tbody>
    </table>
    <form method="post">
      {% for key, value in form_data.items() %}
        {% if value is string %}
          <input type="hidden" name="{{ key }}" value="{{ value }}">
        {% else %}
          {% for item in value %}<input type="hidden" name="{{ key }}" value="{{ item }}">{% endfor %}
        {% endif %}
      {% endfor %}
      <input type="hidden" name="force_create" value="1">
      <div style="display:flex;gap:10px">
        <button type="submit" class="btn btn-danger">Continue Anyway</button>
        <a href="/" class="btn btn-sec">Cancel</a>
      </div>
    </form>
  </div>
</div>
</div>
</main>
</div>
""" + _SB_JS + """
</body></html>"""

# ================================================================
# WORKS LIST
# ================================================================

WORKS_LIST_HTML = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>Works - LabelMind</title>""" + _STYLE + """
</head>
<body>
<div class="app" id="mainApp">
""" + _sidebar("works_list") + """
<main class="main">
""" + _topbar("works") + """
<div class="page">
<div class="ph">
  <div class="ph-left">
    <div class="ph-icon">&#127925;</div>
    <div><div class="ph-title">Works</div><div class="ph-sub">All registered musical works</div></div>
  </div>
  <div class="ph-actions"><a href="/" class="btn btn-primary">+ New Work</a></div>
</div>
<div class="card">
  <div class="card-hd"><div class="card-ico">&#128269;</div><span class="card-title">Search</span></div>
  <div class="card-body">
    <form method="get" style="display:flex;gap:8px">
      <input class="inp" name="q" value="{{ q }}" placeholder="Search work title..." style="max-width:340px">
      <button class="btn btn-sec" type="submit">Search</button>
      {% if q %}<a href="/works" class="btn btn-sec">Clear</a>{% endif %}
    </form>
  </div>
</div>
<div class="card">
  <div class="card-hd"><div class="card-ico">&#128203;</div><span class="card-title">All Works</span></div>
  <div class="tbl-wrap">
    <table class="tbl" style="min-width:960px">
      <thead>
        <tr>
          <th>Work Title</th><th>Session</th><th>Contract Date</th><th>Writers</th>
          <th>Contract</th><th>Signed PDF</th><th>DS Status</th><th>Signed</th>
          <th>Created</th><th></th>
        </tr>
      </thead>
      <tbody>
        {% for work in works %}
        {% set docs = work.contract_documents %}
        {% set first_doc = docs[0] if docs else none %}
        <tr>
          <td style="font-weight:600">{{ work.title }}</td>
          <td style="color:var(--t2)">
            {% if work.batch_id %}
              <a href="/batches/{{ work.batch_id }}" style="color:var(--a)">
                {% if work.batch and work.batch.session_name %}{{ work.batch.session_name }}{% else %}Session #{{ work.batch_id }}{% endif %}
              </a>
            {% else %}--{% endif %}
          </td>
          <td style="color:var(--t2);font-size:12px">{{ work.contract_date.strftime('%b %d, %Y') if work.contract_date else '--' }}</td>
          <td><span style="background:rgba(99,133,255,.1);color:var(--a);border:1px solid rgba(99,133,255,.2);border-radius:99px;padding:2px 8px;font-size:11px;font-weight:700">{{ work.work_writers|length }}</span></td>
          <td>
            {% if docs|length == 1 and first_doc.drive_web_view_link %}
              <a href="{{ first_doc.drive_web_view_link }}" target="_blank" class="file-link" title="{{ first_doc.file_name }}">&#128196; {{ first_doc.file_name | truncate(26,true,'...') }}</a>
            {% elif docs|length > 1 %}
              <a href="/works/{{ work.id }}" class="btn btn-cyan btn-xs">&#128196; {{ docs|length }} docs</a>
            {% else %}--{% endif %}
          </td>
          <td>
            {% set signed_docs = [] %}
            {% for d in docs %}{% if d.signed_pdf_drive_web_view_link %}{% set _ = signed_docs.append(d) %}{% endif %}{% endfor %}
            {% if signed_docs|length == 1 %}
              <a href="{{ signed_docs[0].signed_pdf_drive_web_view_link }}" target="_blank" class="file-link">&#128209; Signed PDF</a>
            {% elif signed_docs|length > 1 %}
              <a href="/works/{{ work.id }}" class="btn btn-success btn-xs">&#128209; {{ signed_docs|length }}</a>
            {% else %}--{% endif %}
          </td>
          <td>
            {% set ns = namespace(ds_st=none) %}
            {% for d in docs %}{% if d.docusign_status and not ns.ds_st %}{% set ns.ds_st = d.docusign_status %}{% endif %}{% endfor %}
            {% if ns.ds_st %}<span class="status s-{{ ns.ds_st }}"><span class="status-dot"></span>{{ ns.ds_st | title }}</span>{% else %}--{% endif %}
          </td>
          <td>
            {% set ns2 = namespace(any_signed=false) %}
            {% for d in docs %}{% if d.status in ['signed','signed_uploaded','signed_complete'] %}{% set ns2.any_signed = true %}{% endif %}{% endfor %}
            {% if ns2.any_signed %}<span class="tag tag-s1">Signed</span>
            {% elif docs %}<span style="color:var(--t3);font-size:11px">Pending</span>
            {% else %}--{% endif %}
          </td>
          <td style="color:var(--t3);font-size:12px">{{ work.created_at.strftime('%b %d, %Y') }}</td>
          <td><a href="/works/{{ work.id }}" class="btn btn-sec btn-sm">View</a></td>
        </tr>
        {% endfor %}
        {% if not works %}<tr class="empty"><td colspan="10">No works found{% if q %} for "{{ q }}"{% endif %}.</td></tr>{% endif %}
      </tbody>
    </table>
  </div>
</div>
</div>
</main>
</div>
""" + _SB_JS + """
</body></html>"""

# ================================================================
# SESSIONS LIST
# ================================================================

BATCHES_LIST_HTML = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>Sessions - LabelMind</title>""" + _STYLE + """
</head>
<body>
<div class="app" id="mainApp">
""" + _sidebar("batches_list") + """
<main class="main">
""" + _topbar("sessions") + """
<div class="page">
<div class="ph">
  <div class="ph-left">
    <div class="ph-icon">&#128230;</div>
    <div><div class="ph-title">Sessions</div><div class="ph-sub">Groups of works ready for contract generation</div></div>
  </div>
  <div class="ph-actions"><a href="/" class="btn btn-primary">+ New Work</a></div>
</div>
<div class="card">
  <div class="card-hd"><div class="card-ico">&#128230;</div><span class="card-title">All Sessions</span></div>
  <div class="tbl-wrap">
    <table class="tbl">
      <thead><tr><th>Session</th><th>Name</th><th>Contract Date</th><th>Status</th><th>Works</th><th>Created</th><th></th></tr></thead>
      <tbody>
        {% for batch in batches %}
        <tr>
          <td style="font-weight:600">Session #{{ batch.id }}</td>
          <td style="color:var(--t2)">{{ batch.session_name or '--' }}</td>
          <td style="color:var(--t2);font-size:12px">{{ batch.contract_date.strftime('%b %d, %Y') }}</td>
          <td><span class="status s-{{ batch.status }}"><span class="status-dot"></span>{{ batch.status | replace('_',' ') | title }}</span></td>
          <td><span style="background:rgba(99,133,255,.1);color:var(--a);border:1px solid rgba(99,133,255,.2);border-radius:99px;padding:2px 8px;font-size:11px;font-weight:700">{{ batch.works|length }}</span></td>
          <td style="color:var(--t3);font-size:12px">{{ batch.created_at.strftime('%b %d, %Y') }}</td>
          <td><a href="/batches/{{ batch.id }}" class="btn btn-sec btn-sm">View</a></td>
        </tr>
        {% endfor %}
        {% if not batches %}<tr class="empty"><td colspan="7">No sessions yet. Create a work to get started.</td></tr>{% endif %}
      </tbody>
    </table>
  </div>
</div>
</div>
</main>
</div>
""" + _SB_JS + """
</body></html>"""

# ================================================================
# SESSION DETAIL
# ================================================================

BATCH_DETAIL_HTML = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>Session {{ batch.id }} - LabelMind</title>""" + _STYLE + """
</head>
<body>
<div class="app" id="mainApp">
""" + _sidebar("batches_list") + """
<main class="main">
""" + _topbar("sessions") + """
<div class="page">
{% with messages = get_flashed_messages() %}{% if messages %}
<div class="flash-list">{% for m in messages %}<div class="flash-item">&#9888; {{ m }}</div>{% endfor %}</div>
{% endif %}{% endwith %}
<div class="ph">
  <div class="ph-left">
    <div class="ph-icon">&#128230;</div>
    <div>
      <div class="ph-title">Session #{{ batch.id }}</div>
      <div class="ph-sub">{{ batch.session_name or 'No name' }} - {{ batch.contract_date.strftime('%b %d, %Y') }}</div>
    </div>
  </div>
  <div class="ph-actions">
    <a href="/batches" class="btn btn-sec btn-sm">Back</a>
    <a href="/?batch_id={{ batch.id }}" class="btn btn-sec btn-sm">+ Add Work</a>
    <form method="post" action="/batches/{{ batch.id }}/generate" id="genForm" style="display:inline">
      <button type="submit" class="btn btn-primary btn-sm" id="genBtn">
        <span id="genLabel">Generate Docs</span><span class="spin" id="genSpin"></span>
      </button>
    </form>
  </div>
</div>
<div class="card">
  <div class="card-hd"><div class="card-ico">&#8505;</div><span class="card-title">Session Info</span></div>
  <div class="card-body">
    <div class="info-grid">
      <div class="info-item"><label>Session Name</label><span>{{ batch.session_name or '--' }}</span></div>
      <div class="info-item"><label>Contract Date</label><span>{{ batch.contract_date.strftime('%B %d, %Y') }}</span></div>
      <div class="info-item"><label>Status</label><span class="status s-{{ batch.status }}"><span class="status-dot"></span>{{ batch.status | replace('_',' ') | title }}</span></div>
      <div class="info-item"><label>Created</label><span>{{ batch.created_at.strftime('%b %d, %Y %H:%M') }}</span></div>
    </div>
  </div>
</div>
<div class="card">
  <div class="card-hd"><div class="card-ico">&#127925;</div><span class="card-title">Works in Session</span></div>
  <div class="tbl-wrap">
    <table class="tbl">
      <thead><tr><th>Work Title</th><th>Writers</th><th>Created</th><th></th></tr></thead>
      <tbody>
        {% for work in works %}
        <tr>
          <td style="font-weight:600">{{ work.title }}</td>
          <td><span style="background:rgba(99,133,255,.1);color:var(--a);border:1px solid rgba(99,133,255,.2);border-radius:99px;padding:2px 8px;font-size:11px;font-weight:700">{{ work.work_writers|length }}</span></td>
          <td style="color:var(--t3);font-size:12px">{{ work.created_at.strftime('%b %d, %Y') }}</td>
          <td><a href="/works/{{ work.id }}" class="btn btn-sec btn-sm">View</a></td>
        </tr>
        {% endfor %}
        {% if not works %}<tr class="empty"><td colspan="4">No works in this session.</td></tr>{% endif %}
      </tbody>
    </table>
  </div>
</div>
<div class="card">
  <div class="card-hd"><div class="card-ico">&#128101;</div><span class="card-title">Writer Summary</span></div>
  <div class="tbl-wrap">
    <table class="tbl">
      <thead><tr><th>Writer</th><th>AKA</th><th>IPI</th><th>PRO</th><th>Works</th><th>Master Contract</th></tr></thead>
      <tbody>
        {% for item in writer_summary %}
        <tr>
          <td style="font-weight:600">{{ item.writer.full_name }}</td>
          <td style="color:var(--t2)">{{ item.writer.writer_aka or '--' }}</td>
          <td style="font-family:var(--fm);font-size:12px;color:var(--t2)">{{ item.writer.ipi or '--' }}</td>
          <td><span class="tag tag-full">{{ item.writer.pro or '--' }}</span></td>
          <td>{{ item.work_count }}</td>
          <td>{% if item.writer.has_master_contract %}<span class="tag tag-s1">Yes</span>{% else %}<span style="color:var(--t3)">No</span>{% endif %}</td>
        </tr>
        {% endfor %}
        {% if not writer_summary %}<tr class="empty"><td colspan="6">No writers in this session.</td></tr>{% endif %}
      </tbody>
    </table>
  </div>
</div>
<div class="card">
  <div class="card-hd"><div class="card-ico">&#128196;</div><span class="card-title">Generated Documents</span></div>
  <div class="tbl-wrap">
    <table class="tbl" style="min-width:860px">
      <thead>
        <tr>
          <th>Writer</th><th>Type</th><th>File</th><th>Generated</th>
          <th>DocuSign</th><th>DS Status</th><th>Certificate</th>
          <th>Upload Signed</th><th>Signed PDF</th><th>Status</th>
        </tr>
      </thead>
      <tbody id="generatedDocumentsBody">
        {% for doc in documents %}
        <tr data-doc-id="{{ doc.id }}">
          <td style="font-weight:600;white-space:nowrap">{{ doc.writer_name_snapshot }}</td>
          <td><span class="tag tag-full">{{ doc.document_type }}</span></td>
          <td>
            {% if doc.drive_web_view_link %}
              <a href="{{ doc.drive_web_view_link }}" target="_blank" class="file-link" title="{{ doc.file_name }}">&#128196; {{ doc.file_name | truncate(30,true,'...') }}</a>
            {% else %}
              <span class="file-link-plain">{{ doc.file_name }}</span>
            {% endif %}
          </td>
          <td style="color:var(--t3);font-size:11.5px">{{ doc.generated_at.strftime('%b %d %H:%M') if doc.generated_at else '--' }}</td>
          <td>
            <form method="post" action="/documents/{{ doc.id }}/send-docusign" class="ds-form">
              <button type="submit" class="btn btn-sec btn-xs ds-btn">
                <span class="ds-lbl">{% if doc.docusign_status == 'completed' %}Resend{% elif doc.docusign_status == 'sent' %}Sent{% elif doc.docusign_status == 'delivered' %}Delivered{% else %}Send{% endif %}</span>
                <span class="spin ds-spin"></span>
              </button>
            </form>
          </td>
          <td>{% if doc.docusign_status %}<span class="status s-{{ doc.docusign_status }}"><span class="status-dot"></span>{{ doc.docusign_status | title }}</span>{% else %}--{% endif %}</td>
          <td>{% if doc.certificate_drive_web_view_link %}<a href="{{ doc.certificate_drive_web_view_link }}" target="_blank" class="btn btn-sec btn-xs">Cert</a>{% else %}--{% endif %}</td>
          <td>
            <form method="post" action="/documents/{{ doc.id }}/upload-signed" enctype="multipart/form-data" class="upl-form">
              <input type="file" name="signed_file" class="upl-inp" required>
              <button type="submit" class="btn btn-success btn-xs">Upload</button>
            </form>
          </td>
          <td>
            {% if doc.signed_pdf_drive_web_view_link %}<a href="{{ doc.signed_pdf_drive_web_view_link }}" target="_blank" class="file-link">&#128209; Signed</a>
            {% elif doc.signed_web_view_link %}<a href="{{ doc.signed_web_view_link }}" target="_blank" class="file-link">&#128209; Signed</a>
            {% else %}--{% endif %}
          </td>
          <td>{% if doc.status %}<span class="status s-{{ doc.status }}"><span class="status-dot"></span>{{ doc.status | replace('_',' ') | title }}</span>{% else %}--{% endif %}</td>
        </tr>
        {% endfor %}
        {% if not documents %}<tr class="empty"><td colspan="10">No documents generated yet. Click Generate Docs above.</td></tr>{% endif %}
      </tbody>
    </table>
  </div>
</div>
</div>
</main>
</div>
""" + _SB_JS + """
<script>
var batchId = {{ batch.id }};

function esc(v) {
  return (v || '').replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;').replace(/"/g,'&quot;');
}

function renderFileCell(doc) {
  if (doc.drive_web_view_link) {
    var short = doc.file_name.length > 30 ? doc.file_name.substring(0, 30) + '...' : doc.file_name;
    return '<a href="' + doc.drive_web_view_link + '" target="_blank" class="file-link" title="' + esc(doc.file_name) + '">&#128196; ' + esc(short) + '</a>';
  }
  return '<span class="file-link-plain">' + esc(doc.file_name) + '</span>';
}

function renderDsBtn(doc) {
  var url = '/documents/' + doc.id + '/send-docusign';
  var lbl = 'Send';
  if (doc.docusign_status === 'completed') lbl = 'Resend';
  else if (doc.docusign_status === 'delivered') lbl = 'Delivered';
  else if (doc.docusign_status === 'sent') lbl = 'Sent';
  return '<form method="post" action="' + url + '" class="ds-form"><button type="submit" class="btn btn-sec btn-xs ds-btn"><span class="ds-lbl">' + lbl + '</span><span class="spin ds-spin"></span></button></form>';
}

function renderStatus(val, cls) {
  if (!val) return '--';
  var c = cls + val.replace(/[ _]/g, '_');
  var l = val.replace(/_/g, ' ').replace(/\b\w/g, function(x) { return x.toUpperCase(); });
  return '<span class="status ' + c + '"><span class="status-dot"></span>' + l + '</span>';
}

function updateDocs(data) {
  var tb = document.getElementById('generatedDocumentsBody');
  if (!tb || !data.documents) return;
  tb.innerHTML = data.documents.map(function(doc) {
    var signedCell = '--';
    if (doc.signed_pdf_drive_web_view_link) {
      signedCell = '<a href="' + doc.signed_pdf_drive_web_view_link + '" target="_blank" class="file-link">&#128209; Signed</a>';
    } else if (doc.signed_web_view_link) {
      signedCell = '<a href="' + doc.signed_web_view_link + '" target="_blank" class="file-link">&#128209; Signed</a>';
    }
    var certCell = doc.certificate_drive_web_view_link
      ? '<a href="' + doc.certificate_drive_web_view_link + '" target="_blank" class="btn btn-sec btn-xs">Cert</a>'
      : '--';
    return '<tr data-doc-id="' + doc.id + '">'
      + '<td style="font-weight:600;white-space:nowrap">' + esc(doc.writer_name_snapshot) + '</td>'
      + '<td><span class="tag tag-full">' + esc(doc.document_type) + '</span></td>'
      + '<td>' + renderFileCell(doc) + '</td>'
      + '<td style="color:var(--t3);font-size:11.5px">' + (doc.generated_at || '--') + '</td>'
      + '<td>' + renderDsBtn(doc) + '</td>'
      + '<td>' + renderStatus(doc.docusign_status, 's-') + '</td>'
      + '<td>' + certCell + '</td>'
      + '<td><form method="post" action="/documents/' + doc.id + '/upload-signed" enctype="multipart/form-data" class="upl-form"><input type="file" name="signed_file" class="upl-inp" required><button type="submit" class="btn btn-success btn-xs">Upload</button></form></td>'
      + '<td>' + signedCell + '</td>'
      + '<td>' + renderStatus(doc.status, 's-') + '</td>'
      + '</tr>';
  }).join('');
  bindDs();
  if (data.documents.length > 0) stopGenSpin();
}

function poll() {
  fetch('/batches/' + batchId + '/status-json', {cache: 'no-store'})
    .then(function(r) { if (r.ok) return r.json(); })
    .then(function(d) { if (d) updateDocs(d); })
    .catch(function(e) { console.error(e); });
}

function stopGenSpin() {
  var btn = document.getElementById('genBtn');
  if (!btn) return;
  btn.disabled = false;
  document.getElementById('genSpin').classList.remove('on');
  document.getElementById('genLabel').textContent = 'Generate Docs';
}

function bindDs() {
  document.querySelectorAll('.ds-form').forEach(function(f) {
    if (f.dataset.bound) return;
    f.dataset.bound = '1';
    f.addEventListener('submit', function(e) {
      e.preventDefault();
      var btn = f.querySelector('.ds-btn');
      var spin = f.querySelector('.ds-spin');
      var lbl = f.querySelector('.ds-lbl');
      if (btn) btn.disabled = true;
      if (spin) spin.classList.add('on');
      if (lbl) lbl.textContent = 'Sending...';
      setTimeout(function() { f.submit(); }, 150);
    });
  });
}

document.addEventListener('DOMContentLoaded', function() {
  bindDs();
  var gf = document.getElementById('genForm');
  if (gf) {
    gf.addEventListener('submit', function(e) {
      e.preventDefault();
      var btn = document.getElementById('genBtn');
      btn.disabled = true;
      document.getElementById('genSpin').classList.add('on');
      document.getElementById('genLabel').textContent = 'Generating...';
      setTimeout(function() { gf.submit(); }, 150);
      setTimeout(function() { window.location.reload(); }, 5000);
    });
  }
  setInterval(poll, 5000);
});
</script>
</body></html>"""

# ================================================================
# WORK DETAIL
# ================================================================

WORK_DETAIL_HTML = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>{{ work.title }} - LabelMind</title>""" + _STYLE + """
</head>
<body>
<div class="app" id="mainApp">
""" + _sidebar("works_list") + """
<main class="main">
""" + _topbar("works") + """
<div class="page">
<div class="ph">
  <div class="ph-left">
    <div class="ph-icon">&#127925;</div>
    <div>
      <div class="ph-title">{{ work.title }}</div>
      <div class="ph-sub">{{ work.batch.session_name if work.batch and work.batch.session_name else 'No session' }} - {{ work.contract_date.strftime('%b %d, %Y') if work.contract_date else '--' }}</div>
    </div>
  </div>
  <div class="ph-actions">
    {% if work.batch_id %}<a href="/batches/{{ work.batch_id }}" class="btn btn-sec btn-sm">View Session</a>{% endif %}
    <a href="/works" class="btn btn-sec btn-sm">Back</a>
  </div>
</div>
<div class="card">
  <div class="card-hd"><div class="card-ico">&#128203;</div><span class="card-title">Work Info</span></div>
  <div class="card-body">
    <div class="info-grid">
      <div class="info-item"><label>Session Name</label><span>{{ work.batch.session_name if work.batch and work.batch.session_name else '--' }}</span></div>
      <div class="info-item"><label>Session</label>{% if work.batch_id %}<a href="/batches/{{ work.batch_id }}">Session #{{ work.batch_id }}</a>{% else %}<span>--</span>{% endif %}</div>
      <div class="info-item"><label>Contract Date</label><span>{{ work.contract_date.strftime('%B %d, %Y') if work.contract_date else '--' }}</span></div>
      <div class="info-item"><label>Created</label><span>{{ work.created_at.strftime('%b %d, %Y %H:%M') }}</span></div>
    </div>
  </div>
</div>
<div class="card">
  <div class="card-hd"><div class="card-ico">&#128101;</div><span class="card-title">Writers &amp; Splits</span></div>
  <div class="tbl-wrap">
    <table class="tbl">
      <thead><tr><th>Writer</th><th>AKA</th><th>IPI</th><th>PRO</th><th>Split %</th><th>Publisher</th><th>Pub IPI</th><th>Master</th></tr></thead>
      <tbody>
        {% for ww in work.work_writers %}
        <tr>
          <td style="font-weight:600">{{ ww.writer.full_name }}</td>
          <td style="color:var(--t2)">{{ ww.writer.writer_aka or '--' }}</td>
          <td style="font-family:var(--fm);font-size:12px;color:var(--t2)">{{ ww.writer.ipi or '--' }}</td>
          <td><span class="tag tag-full">{{ ww.writer.pro or '--' }}</span></td>
          <td><span style="font-family:var(--fm);font-size:13px;font-weight:600;color:var(--a)">{{ "%.2f"|format(ww.writer_percentage) }}%</span></td>
          <td style="color:var(--t2)">{{ ww.publisher or '--' }}</td>
          <td style="font-family:var(--fm);font-size:12px;color:var(--t2)">{{ ww.publisher_ipi or '--' }}</td>
          <td>{% if ww.writer.has_master_contract %}<span class="tag tag-s1">Yes</span>{% else %}<span style="color:var(--t3)">No</span>{% endif %}</td>
        </tr>
        {% endfor %}
      </tbody>
    </table>
  </div>
</div>
<div class="card">
  <div class="card-hd"><div class="card-ico">&#128196;</div><span class="card-title">Generated Documents</span></div>
  <div class="tbl-wrap">
    <table class="tbl" style="min-width:860px">
      <thead>
        <tr>
          <th>Writer</th><th>Type</th><th>File</th><th>Generated At</th>
          <th>DocuSign</th><th>DS Status</th><th>Certificate</th><th>Signed PDF</th><th>Status</th>
        </tr>
      </thead>
      <tbody>
        {% for doc in documents %}
        <tr data-doc-id="{{ doc.id }}">
          <td style="font-weight:600;white-space:nowrap">{{ doc.writer_name_snapshot }}</td>
          <td><span class="tag tag-full">{{ doc.document_type }}</span></td>
          <td>
            {% if doc.drive_web_view_link %}
              <a href="{{ doc.drive_web_view_link }}" target="_blank" class="file-link" title="{{ doc.file_name }}">&#128196; {{ doc.file_name | truncate(30,true,'...') }}</a>
            {% else %}
              <span class="file-link-plain">{{ doc.file_name }}</span>
            {% endif %}
          </td>
          <td style="color:var(--t3);font-size:11.5px">{{ doc.generated_at.strftime('%b %d, %Y') if doc.generated_at else '--' }}</td>
          <td>
            <form method="post" action="/documents/{{ doc.id }}/send-docusign" class="ds-form">
              <button type="submit" class="btn btn-sec btn-xs ds-btn">
                <span class="ds-lbl">{% if doc.docusign_status == 'completed' %}Resend{% elif doc.docusign_status == 'sent' %}Sent{% elif doc.docusign_status == 'delivered' %}Delivered{% else %}Send{% endif %}</span>
                <span class="spin ds-spin"></span>
              </button>
            </form>
          </td>
          <td>{% if doc.docusign_status %}<span class="status s-{{ doc.docusign_status }}"><span class="status-dot"></span>{{ doc.docusign_status | title }}</span>{% else %}--{% endif %}</td>
          <td>{% if doc.certificate_drive_web_view_link %}<a href="{{ doc.certificate_drive_web_view_link }}" target="_blank" class="btn btn-sec btn-xs">Cert</a>{% else %}--{% endif %}</td>
          <td>
            {% if doc.signed_pdf_drive_web_view_link %}<a href="{{ doc.signed_pdf_drive_web_view_link }}" target="_blank" class="file-link">&#128209; Signed</a>
            {% elif doc.signed_web_view_link %}<a href="{{ doc.signed_web_view_link }}" target="_blank" class="file-link">&#128209; Signed</a>
            {% else %}--{% endif %}
          </td>
          <td>{% if doc.status %}<span class="status s-{{ doc.status }}"><span class="status-dot"></span>{{ doc.status | replace('_',' ') | title }}</span>{% else %}--{% endif %}</td>
        </tr>
        {% endfor %}
        {% if not documents %}<tr class="empty"><td colspan="9">No documents generated yet.</td></tr>{% endif %}
      </tbody>
    </table>
  </div>
</div>
</div>
</main>
</div>
""" + _SB_JS + """
<script>
document.querySelectorAll('.ds-form').forEach(function(f) {
  f.addEventListener('submit', function(e) {
    e.preventDefault();
    var btn = f.querySelector('.ds-btn');
    var spin = f.querySelector('.ds-spin');
    var lbl = f.querySelector('.ds-lbl');
    if (btn) btn.disabled = true;
    if (spin) spin.classList.add('on');
    if (lbl) lbl.textContent = 'Sending...';
    setTimeout(function() { f.submit(); }, 150);
  });
});
</script>
</body></html>"""


# ================================================================
# HELPERS
# ================================================================

def auth_required():
    if not (TEAM_USERNAME and TEAM_PASSWORD):
        return False
    return not session.get("logged_in")


def get_or_create_camp(existing_camp_id, new_camp_name):
    new_camp_name = (new_camp_name or "").strip()
    if new_camp_name:
        existing = Camp.query.filter(func.lower(Camp.name) == new_camp_name.lower()).first()
        if existing:
            return existing
        camp = Camp(name=new_camp_name)
        db.session.add(camp)
        db.session.flush()
        return camp
    if existing_camp_id:
        return Camp.query.get(int(existing_camp_id))
    return None


def find_existing_writer(selected_writer_id):
    if selected_writer_id:
        writer = Writer.query.get(int(selected_writer_id))
        if writer:
            return writer
    return None


def render_docx_template(template_path, data, works_for_table=None):
    if not os.path.exists(template_path):
        raise FileNotFoundError(template_path)
    doc = Document(template_path)

    def replace_all(paragraph):
        text = "".join(run.text for run in paragraph.runs)
        for k, v in data.items():
            text = text.replace("[[" + k + "]]", str(v))
        for run in paragraph.runs:
            run.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = text

    for p in doc.paragraphs:
        replace_all(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_all(p)
    for p in doc.paragraphs:
        if "[[ContractTable]]" in p.text:
            table = doc.add_table(rows=1, cols=5)
            table.style = "Table Grid"
            table.rows[0].cells[0].text = "Work Title"
            table.rows[0].cells[1].text = "Songwriter Name"
            table.rows[0].cells[2].text = "Songwriter Share"
            table.rows[0].cells[3].text = "Publisher Name"
            table.rows[0].cells[4].text = "Publisher Share"
            for item in works_for_table or []:
                row = table.add_row().cells
                row[0].text = item.get("work_title", "")
                row[1].text = item.get("writer_name", "")
                row[2].text = item.get("writer_percentage", "")
                row[3].text = item.get("publisher", "")
                row[4].text = item.get("writer_percentage", "")
            p.text = ""
            p._element.addnext(table._element)
            break

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def collect_form_context():
    selected_batch_id = (
        request.form.get("existing_batch_id")
        or request.args.get("batch_id")
        or ""
    )
    return {
        "batches": GenerationBatch.query.order_by(GenerationBatch.created_at.desc()).all(),
        "default_publisher_address": DEFAULT_PUBLISHER_ADDRESS,
        "default_publisher_city": DEFAULT_PUBLISHER_CITY,
        "default_publisher_state": DEFAULT_PUBLISHER_STATE,
        "default_publisher_zip": DEFAULT_PUBLISHER_ZIP,
        "force_create": request.form.get("force_create", ""),
        "selected_batch_id": selected_batch_id,
    }


def get_batch_writer_summary(batch_id):
    work_writers = WorkWriter.query.join(Work).filter(Work.batch_id == batch_id).all()
    grouped = {}
    for ww in work_writers:
        if ww.writer_id not in grouped:
            grouped[ww.writer_id] = {"writer": ww.writer, "work_titles": set()}
        grouped[ww.writer_id]["work_titles"].add(ww.work.title)
    summary = []
    for item in grouped.values():
        summary.append({"writer": item["writer"], "work_count": len(item["work_titles"])})
    summary.sort(key=lambda x: x["writer"].full_name.lower())
    return summary


# ================================================================
# ROUTES
# ================================================================

@app.route("/login", methods=["GET", "POST"])
def login():
    if not (TEAM_USERNAME and TEAM_PASSWORD):
        return redirect(url_for("formulario"))
    if request.method == "POST":
        username = request.form.get("username", "")
        password = request.form.get("password", "")
        if username == TEAM_USERNAME and password == TEAM_PASSWORD:
            session["logged_in"] = True
            return redirect(url_for("formulario"))
        flash("Incorrect username or password.")
    return render_template_string(LOGIN_HTML)


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))


@app.route("/", methods=["GET", "POST"])
def formulario():
    if auth_required():
        return redirect(url_for("login"))

    if request.method == "POST":
        work_title = (request.form.get("work_title") or "").strip()
        contract_date_str = (request.form.get("contract_date") or "").strip()

        if not work_title:
            flash("Work title is required.")
            return render_template_string(FORM_HTML, **collect_form_context())
        if not contract_date_str:
            flash("Contract date is required.")
            return render_template_string(FORM_HTML, **collect_form_context())

        try:
            contract_date = datetime.datetime.strptime(contract_date_str, "%Y-%m-%d").date()
        except ValueError:
            flash("Please enter a valid contract date.")
            return render_template_string(FORM_HTML, **collect_form_context())

        writer_ids = request.form.getlist("writer_id")
        first_names = request.form.getlist("writer_first_name")
        middle_names = request.form.getlist("writer_middle_name")
        last_names_list = request.form.getlist("writer_last_names")
        writer_akas = request.form.getlist("writer_aka")
        ipis = request.form.getlist("writer_ipi")
        emails = request.form.getlist("writer_email")
        pros = request.form.getlist("writer_pro")
        percentages = request.form.getlist("writer_percentage")
        publishers = request.form.getlist("writer_publisher")
        publisher_ipis = request.form.getlist("publisher_ipi")
        publisher_addresses = request.form.getlist("publisher_address")
        publisher_cities = request.form.getlist("publisher_city")
        publisher_states = request.form.getlist("publisher_state")
        publisher_zips = request.form.getlist("publisher_zip_code")
        addresses = request.form.getlist("writer_address")
        cities = request.form.getlist("writer_city")
        states = request.form.getlist("writer_state")
        zip_codes = request.form.getlist("writer_zip_code")

        writer_rows = []
        total_split = 0.0

        for i in range(len(first_names)):
            first_name = (first_names[i] or "").strip()
            middle_name = (middle_names[i] or "").strip()
            last_names = (last_names_list[i] or "").strip()
            full_name = build_full_name(first_name, middle_name, last_names)
            if not full_name:
                continue
            split_value = parse_float(percentages[i] if i < len(percentages) else "0")
            if split_value <= 0:
                flash("Writer '" + full_name + "' must have a split greater than 0.")
                return render_template_string(FORM_HTML, **collect_form_context())
            total_split += split_value
            writer_rows.append({
                "selected_writer_id": writer_ids[i] if i < len(writer_ids) else "",
                "first_name": first_name,
                "middle_name": middle_name,
                "last_names": last_names,
                "full_name": full_name,
                "writer_aka": (writer_akas[i] or "").strip(),
                "ipi": (ipis[i] or "").strip(),
                "email": (emails[i] or "").strip(),
                "pro": (pros[i] or "").strip(),
                "writer_percentage": split_value,
                "publisher": (publishers[i] or "").strip(),
                "publisher_ipi": (publisher_ipis[i] or "").strip(),
                "publisher_address": (publisher_addresses[i] or DEFAULT_PUBLISHER_ADDRESS).strip(),
                "publisher_city": (publisher_cities[i] or DEFAULT_PUBLISHER_CITY).strip(),
                "publisher_state": (publisher_states[i] or DEFAULT_PUBLISHER_STATE).strip(),
                "publisher_zip_code": (publisher_zips[i] or DEFAULT_PUBLISHER_ZIP).strip(),
                "address": (addresses[i] or "").strip(),
                "city": (cities[i] or "").strip(),
                "state": (states[i] or "").strip(),
                "zip_code": (zip_codes[i] or "").strip(),
            })

        if not writer_rows:
            flash("Add at least one writer.")
            return render_template_string(FORM_HTML, **collect_form_context())

        if abs(total_split - 100.0) >= 0.001:
            flash("Total writer split must equal 100%. Current total: " + str(round(total_split, 2)) + "%")
            return render_template_string(FORM_HTML, **collect_form_context())

        seen_writer_ids = set()
        seen_ipis = set()
        seen_names = set()

        for row in writer_rows:
            selected_writer_id = (row["selected_writer_id"] or "").strip()
            ipi = (row["ipi"] or "").strip()
            normalized_name = normalize_text(row["full_name"])
            if selected_writer_id:
                if selected_writer_id in seen_writer_ids:
                    flash("Duplicate writer selected in this work: " + row["full_name"])
                    return render_template_string(FORM_HTML, **collect_form_context())
                seen_writer_ids.add(selected_writer_id)
            if ipi:
                ipi_key = ipi.lower()
                if ipi_key in seen_ipis:
                    flash("Duplicate IPI in this work: " + ipi)
                    return render_template_string(FORM_HTML, **collect_form_context())
                seen_ipis.add(ipi_key)
            else:
                if normalized_name in seen_names:
                    flash("Duplicate writer name in this work: " + row["full_name"])
                    return render_template_string(FORM_HTML, **collect_form_context())
                seen_names.add(normalized_name)

        for row in writer_rows:
            if row["ipi"]:
                existing_ipi_writer = Writer.query.filter(func.lower(Writer.ipi) == row["ipi"].lower()).first()
                if existing_ipi_writer:
                    selected_id = (row["selected_writer_id"] or "").strip()
                    if not selected_id or str(existing_ipi_writer.id) != selected_id:
                        flash("IPI " + row["ipi"] + " already belongs to " + existing_ipi_writer.full_name + ". Please select the existing writer.")
                        return render_template_string(FORM_HTML, **collect_form_context())

        warnings = []
        for row in writer_rows:
            if not row["ipi"]:
                existing_name_writer = Writer.query.filter(
                    func.lower(Writer.full_name) == normalize_text(row["full_name"])
                ).first()
                if existing_name_writer:
                    warnings.append("Writer '" + row["full_name"] + "' already exists in the system without using an IPI match.")

        normalized_title = normalize_title(work_title)
        writer_identity_set = sorted([build_writer_identity_from_row(row) for row in writer_rows])

        possible_duplicates = []
        existing_works = Work.query.filter_by(normalized_title=normalized_title).all()
        for existing_work in existing_works:
            existing_identities = sorted([
                build_writer_identity_from_workwriter(ww) for ww in existing_work.work_writers
            ])
            if existing_identities == writer_identity_set:
                possible_duplicates.append({
                    "title": existing_work.title,
                    "camp_name": existing_work.batch.session_name if existing_work.batch and existing_work.batch.session_name else "",
                    "created_at": existing_work.created_at.strftime("%Y-%m-%d"),
                })

        if possible_duplicates and not request.form.get("force_create"):
            form_data = {}
            for key in request.form.keys():
                values = request.form.getlist(key)
                form_data[key] = values if len(values) > 1 else values[0]
            return render_template_string(
                DUPLICATE_WARNING_HTML,
                duplicates=possible_duplicates,
                form_data=form_data,
            )

        for warning in warnings:
            flash(warning)

        existing_batch_id = (request.form.get("existing_batch_id") or "").strip()
        new_session_name = (request.form.get("new_session_name") or "").strip()

        if existing_batch_id:
            batch = GenerationBatch.query.get(int(existing_batch_id))
            if not batch:
                flash("Selected session was not found.")
                return render_template_string(FORM_HTML, **collect_form_context())
            contract_date = batch.contract_date
        else:
            if not new_session_name:
                flash("Please enter a new session name or select an existing session.")
                return render_template_string(FORM_HTML, **collect_form_context())

            batch = GenerationBatch(
                session_name=build_session_name(new_session_name),
                contract_date=contract_date,
                created_by="",
                status="draft",
            )
            db.session.add(batch)
            db.session.flush()

        work = Work(
            title=work_title,
            normalized_title=normalized_title,
            camp_id=None,
            batch_id=batch.id,
            contract_date=contract_date,
        )
        db.session.add(work)
        db.session.flush()

        for row in writer_rows:
            writer = find_existing_writer(row["selected_writer_id"])

            if not writer and row["ipi"]:
                writer = Writer.query.filter(
                    func.lower(Writer.ipi) == row["ipi"].lower()
                ).first()

            if not writer and row["full_name"]:
                writer = Writer.query.filter(
                    func.lower(Writer.full_name) == row["full_name"].lower()
                ).first()

            if writer:
                writer.first_name = row["first_name"] or writer.first_name
                writer.middle_name = row["middle_name"] or writer.middle_name
                writer.last_names = row["last_names"] or writer.last_names
                writer.full_name = row["full_name"] or writer.full_name
                writer.writer_aka = row["writer_aka"] or writer.writer_aka
                writer.ipi = row["ipi"] or writer.ipi
                writer.email = row["email"] or writer.email
                writer.pro = row["pro"] or writer.pro
                writer.address = row["address"] or writer.address
                writer.city = row["city"] or writer.city
                writer.state = row["state"] or writer.state
                writer.zip_code = row["zip_code"] or writer.zip_code
            else:
                writer = Writer(
                    first_name=row["first_name"],
                    middle_name=row["middle_name"],
                    last_names=row["last_names"],
                    full_name=row["full_name"],
                    writer_aka=row["writer_aka"],
                    ipi=row["ipi"] or None,
                    email=row["email"],
                    pro=row["pro"],
                    address=row["address"],
                    city=row["city"],
                    state=row["state"],
                    zip_code=row["zip_code"],
                    has_master_contract=False,
                )
                db.session.add(writer)
                db.session.flush()

            work_writer = WorkWriter(
                work_id=work.id,
                writer_id=writer.id,
                writer_percentage=row["writer_percentage"],
                publisher=row["publisher"],
                publisher_ipi=row["publisher_ipi"],
                publisher_address=row["publisher_address"],
                publisher_city=row["publisher_city"],
                publisher_state=row["publisher_state"],
                publisher_zip_code=row["publisher_zip_code"],
            )
            db.session.add(work_writer)

        db.session.commit()
        return redirect(url_for("batch_detail", batch_id=batch.id))

    return render_template_string(FORM_HTML, **collect_form_context())


@app.route("/writers/search")
def search_writers():
    if auth_required():
        return jsonify([])
    q = (request.args.get("q") or "").strip()
    if len(q) < 2:
        return jsonify([])
    like_q = "%" + q.lower() + "%"
    writers = (
        Writer.query
        .filter(or_(
            func.lower(Writer.full_name).like(like_q),
            func.lower(Writer.first_name).like(like_q),
            func.lower(Writer.middle_name).like(like_q),
            func.lower(Writer.last_names).like(like_q),
            func.lower(Writer.writer_aka).like(like_q),
            func.lower(Writer.ipi).like(like_q),
        ))
        .order_by(Writer.full_name.asc())
        .limit(8)
        .all()
    )
    return jsonify([
        {
            "id": w.id,
            "first_name": w.first_name,
            "middle_name": w.middle_name,
            "last_names": w.last_names,
            "full_name": w.full_name,
            "writer_aka": w.writer_aka,
            "ipi": w.ipi or "",
            "email": w.email or "",
            "pro": w.pro,
            "address": w.address,
            "city": w.city,
            "state": w.state,
            "zip_code": w.zip_code,
            "has_master_contract": w.has_master_contract,
            "default_publisher": default_publisher_for_pro(w.pro),
            "default_publisher_ipi": default_publisher_ipi_for_pro(w.pro),
        }
        for w in writers
    ])


@app.route("/works")
def works_list():
    if auth_required():
        return redirect(url_for("login"))
    q = (request.args.get("q") or "").strip()
    query = Work.query
    if q:
        query = query.filter(func.lower(Work.title).like("%" + q.lower() + "%"))
    works = query.order_by(Work.created_at.desc()).all()
    return render_template_string(WORKS_LIST_HTML, works=works, q=q)


@app.route("/batches")
def batches_list():
    if auth_required():
        return redirect(url_for("login"))
    batches = GenerationBatch.query.order_by(GenerationBatch.created_at.desc()).all()
    return render_template_string(BATCHES_LIST_HTML, batches=batches)


@app.route("/batches/<int:batch_id>")
def batch_detail(batch_id):
    if auth_required():
        return redirect(url_for("login"))
    batch = GenerationBatch.query.get_or_404(batch_id)
    works = Work.query.filter_by(batch_id=batch.id).order_by(Work.created_at.asc()).all()
    documents = ContractDocument.query.filter_by(batch_id=batch.id).order_by(ContractDocument.generated_at.desc()).all()
    writer_summary = get_batch_writer_summary(batch.id)
    return render_template_string(BATCH_DETAIL_HTML, batch=batch, works=works, documents=documents, writer_summary=writer_summary)


@app.route("/batches/<int:batch_id>/status-json")
def batch_status_json(batch_id):
    if auth_required():
        return jsonify({"error": "unauthorized"}), 401
    batch = GenerationBatch.query.get_or_404(batch_id)
    documents = ContractDocument.query.filter_by(batch_id=batch.id).order_by(ContractDocument.generated_at.asc()).all()
    return jsonify({
        "batch_id": batch.id,
        "status": batch.status,
        "documents": [
            {
                "id": doc.id,
                "writer_name_snapshot": doc.writer_name_snapshot,
                "document_type": doc.document_type,
                "file_name": doc.file_name,
                "generated_at": doc.generated_at.strftime("%b %d %H:%M") if doc.generated_at else "",
                "drive_web_view_link": doc.drive_web_view_link,
                "docusign_status": doc.docusign_status,
                "status": doc.status,
                "signed_pdf_drive_web_view_link": getattr(doc, "signed_pdf_drive_web_view_link", None),
                "signed_web_view_link": getattr(doc, "signed_web_view_link", None),
                "certificate_drive_web_view_link": getattr(doc, "certificate_drive_web_view_link", None),
            }
            for doc in documents
        ],
    })


@app.route("/batches/<int:batch_id>/generate", methods=["POST"])
def generate_batch_documents(batch_id):
    if auth_required():
        return redirect(url_for("login"))

    batch = GenerationBatch.query.get_or_404(batch_id)
    work_writers = (
        WorkWriter.query.join(Work)
        .filter(Work.batch_id == batch.id)
        .order_by(Work.id.asc(), WorkWriter.id.asc())
        .all()
    )

    if not work_writers:
        flash("No works found in this session.")
        return redirect(url_for("batch_detail", batch_id=batch.id))

    grouped = {}
    for ww in work_writers:
        if ww.writer_id not in grouped:
            grouped[ww.writer_id] = {"writer": ww.writer, "rows": []}
        grouped[ww.writer_id]["rows"].append(ww)

    existing_docs = ContractDocument.query.filter_by(batch_id=batch.id).all()
    for doc in existing_docs:
        db.session.delete(doc)
    db.session.flush()

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
        for writer_id, item in grouped.items():
            writer = item["writer"]
            rows = item["rows"]

            works_for_table = []
            for ww in rows:
                works_for_table.append({
                    "work_title": ww.work.title,
                    "writer_name": writer.full_name,
                    "writer_percentage": str(round(ww.writer_percentage, 2)) + "%",
                    "publisher": ww.publisher or "",
                })

            first_work = rows[0].work
            first_work_writer = rows[0]

            if writer.has_master_contract:
                document_type = "schedule_1"
                template_path = SCHEDULE_1_TEMPLATE
                prefix = "S1"
            else:
                document_type = "full_contract"
                template_path = FULL_CONTRACT_TEMPLATE
                prefix = "FULL"

            contract_date = batch.contract_date
            day = contract_date.day
            suffix = "th" if 11 <= day <= 13 else {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")

            data = {
                "Date": contract_date.strftime("%B") + " " + str(day) + suffix + ", " + str(contract_date.year),
                "Fecha": format_date(contract_date, format="d 'de' MMMM 'del' y", locale="es"),
                "WriterName": writer.full_name,
                "WriterFirstName": writer.first_name,
                "WriterMiddleName": writer.middle_name,
                "WriterLastNames": writer.last_names,
                "WriterAKA": writer.writer_aka or "",
                "WriterIPI": writer.ipi or "",
                "WriterAddress": writer.address or "",
                "WriterCity": writer.city or "",
                "WriterState": writer.state or "",
                "WriterZipCode": writer.zip_code or "",
                "PRO": writer.pro or "",
                "PublisherName": first_work_writer.publisher or "",
                "PublisherIPI": first_work_writer.publisher_ipi or "",
                "PublisherAddress": first_work_writer.publisher_address or "",
                "PublisherCity": first_work_writer.publisher_city or "",
                "PublisherState": first_work_writer.publisher_state or "",
                "PublisherZipCode": first_work_writer.publisher_zip_code or "",
                "WorkTitle": first_work.title,
            }

            file_buffer = render_docx_template(template_path, data, works_for_table=works_for_table)
            file_bytes = file_buffer.getvalue()

            batch_label = batch.camp.name if batch.camp else "batch_" + str(batch.id)
            file_name = prefix + "_" + slugify(writer.full_name) + "_" + slugify(batch_label) + "_" + batch.contract_date.isoformat() + ".docx"

            zip_file.writestr(file_name, file_bytes)

            drive_info = {"file_id": None, "web_view_link": None}
            if GOOGLE_DRIVE_FOLDER_ID and GOOGLE_SERVICE_ACCOUNT_JSON:
                try:
                    drive_info = upload_bytes_to_drive(
                        file_name=file_name,
                        file_bytes=file_bytes,
                        parent_folder_id=GOOGLE_DRIVE_FOLDER_ID,
                        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                    app.logger.warning("DRIVE SUCCESS: %s", drive_info)
                except Exception as e:
                    app.logger.error("DRIVE FAILURE: %s", e)
                    traceback.print_exc()
                    flash("Drive upload failed for " + file_name + ": " + str(e))
            else:
                flash("Drive upload skipped: missing GOOGLE_DRIVE_FOLDER_ID or GOOGLE_SERVICE_ACCOUNT_JSON")

            doc_record = ContractDocument(
                batch_id=batch.id,
                work_id=first_work.id,
                writer_id=writer.id,
                document_type=document_type,
                file_name=file_name,
                writer_name_snapshot=writer.full_name,
                work_title_snapshot=", ".join(sorted({ww.work.title for ww in rows})),
                drive_file_id=drive_info["file_id"],
                drive_web_view_link=drive_info["web_view_link"],
                status="generated",
            )
            db.session.add(doc_record)

            if document_type == "full_contract":
                writer.has_master_contract = True

    batch.status = "docs_generated"
    db.session.commit()

    zip_buffer.seek(0)
    return send_file(
        zip_buffer,
        as_attachment=True,
        download_name="session_" + str(batch.id) + "_documents.zip",
        mimetype="application/zip",
    )


@app.route("/docusign/webhook", methods=["POST"])
def docusign_webhook():
    raw_data = request.data
    app.logger.warning("RAW BODY: %s", raw_data)
    try:
        root = ET.fromstring(raw_data)
        envelope_id = None
        raw_status = ""
        for elem in root.iter():
            tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
            text = (elem.text or "").strip()
            if tag == "EnvelopeID" and text:
                envelope_id = text
            if tag == "Status" and text:
                raw_status = text.lower()

        if not envelope_id:
            return "ok", 200

        document = ContractDocument.query.filter_by(docusign_envelope_id=envelope_id).first()
        if not document:
            return "ok", 200

        if "completed" in raw_status:
            normalized_status = "completed"
        elif "delivered" in raw_status:
            normalized_status = "delivered"
        elif "sent" in raw_status:
            normalized_status = "sent"
        elif "declined" in raw_status:
            normalized_status = "declined"
        elif "voided" in raw_status:
            normalized_status = "voided"
        else:
            normalized_status = raw_status or document.docusign_status or "sent"

        document.docusign_status = normalized_status

        if normalized_status == "completed":
            api_client = get_docusign_api_client()
            envelopes_api = EnvelopesApi(api_client)

            signed_bytes = envelopes_api.get_document(
                account_id=DOCUSIGN_ACCOUNT_ID,
                envelope_id=document.docusign_envelope_id,
                document_id="combined",
            )
            signed_file_name = document.file_name.rsplit(".", 1)[0] + "_SIGNED.pdf"
            signed_drive_info = upload_bytes_to_drive(
                file_name=signed_file_name,
                file_bytes=signed_bytes,
                parent_folder_id=GOOGLE_DRIVE_FOLDER_ID,
                mime_type="application/pdf",
            )
            document.signed_pdf_drive_file_id = signed_drive_info.get("file_id")
            document.signed_pdf_drive_web_view_link = signed_drive_info.get("web_view_link")

            certificate_bytes = envelopes_api.get_document(
                account_id=DOCUSIGN_ACCOUNT_ID,
                envelope_id=document.docusign_envelope_id,
                document_id="certificate",
            )
            certificate_file_name = document.file_name.rsplit(".", 1)[0] + "_CERTIFICATE.pdf"
            certificate_drive_info = upload_bytes_to_drive(
                file_name=certificate_file_name,
                file_bytes=certificate_bytes,
                parent_folder_id=GOOGLE_DRIVE_FOLDER_ID,
                mime_type="application/pdf",
            )
            document.certificate_drive_file_id = certificate_drive_info.get("file_id")
            document.certificate_drive_web_view_link = certificate_drive_info.get("web_view_link")
            document.completed_at = datetime.datetime.utcnow()
            document.status = "signed"

        db.session.commit()
    except Exception as e:
        db.session.rollback()
        app.logger.error("DocuSign webhook failed: %s", e)
    return "ok", 200


@app.route("/documents/<int:document_id>/send-docusign", methods=["POST"])
def send_document_docusign(document_id):
    document = None
    try:
        if auth_required():
            return redirect(url_for("login"))

        document = ContractDocument.query.get(document_id)
        if not document:
            flash("This document no longer exists. Please refresh the session page.")
            return redirect(request.referrer or url_for("batches_list"))

        writer = Writer.query.get(document.writer_id)

        if not writer:
            flash("Writer not found.")
            return redirect(url_for("batch_detail", batch_id=document.batch_id))
        if not getattr(writer, "email", None):
            flash("Writer email is required before sending to DocuSign.")
            return redirect(url_for("batch_detail", batch_id=document.batch_id))
        if not document.drive_file_id:
            flash("Generated document file is missing.")
            return redirect(url_for("batch_detail", batch_id=document.batch_id))
        if not DOCUSIGN_ACCOUNT_ID or not DOCUSIGN_INTEGRATION_KEY or not DOCUSIGN_USER_ID or not DOCUSIGN_PRIVATE_KEY:
            flash("DocuSign environment variables are not fully configured.")
            return redirect(url_for("batch_detail", batch_id=document.batch_id))

        service = get_drive_service()
        file_bytes = service.files().get_media(fileId=document.drive_file_id, supportsAllDrives=True).execute()

        api_client = get_docusign_api_client()
        envelopes_api = EnvelopesApi(api_client)

        doc_b64 = base64.b64encode(file_bytes).decode("ascii")
        ds_document = DocusignDocument(
            document_base64=doc_b64,
            name=document.file_name,
            file_extension="docx",
            document_id="1",
        )

        signer = Signer(
            email=writer.email,
            name=writer.full_name,
            recipient_id="1",
            routing_order="1",
        )
        sign_here = SignHere(
            anchor_string="[[DS_SIGN_HERE]]",
            anchor_units="pixels",
            anchor_x_offset="0",
            anchor_y_offset="0",
        )
        signer.tabs = Tabs(sign_here_tabs=[sign_here])

        webhook_url = request.url_root.rstrip("/") + url_for("docusign_webhook")
        event_notification = {
            "url": webhook_url,
            "loggingEnabled": "true",
            "requireAcknowledgment": "true",
            "includeEnvelopeVoidReason": "true",
            "includeTimeZone": "true",
            "includeSenderAccountAsCustomField": "true",
            "envelopeEvents": [
                {"envelopeEventStatusCode": "sent"},
                {"envelopeEventStatusCode": "delivered"},
                {"envelopeEventStatusCode": "completed"},
                {"envelopeEventStatusCode": "declined"},
                {"envelopeEventStatusCode": "voided"},
            ],
        }

        envelope_definition = EnvelopeDefinition(
            email_subject="Please sign: " + document.file_name,
            documents=[ds_document],
            recipients=Recipients(signers=[signer]),
            status="sent",
            event_notification=event_notification,
        )

        result = envelopes_api.create_envelope(
            account_id=DOCUSIGN_ACCOUNT_ID,
            envelope_definition=envelope_definition,
        )

        document.docusign_envelope_id = result.envelope_id
        document.docusign_status = "sent"
        document.sent_for_signature_at = datetime.datetime.utcnow()
        db.session.commit()

        flash("Sent to DocuSign successfully.")
        return redirect(url_for("batch_detail", batch_id=document.batch_id))

    except Exception as e:
        db.session.rollback()
        app.logger.error("DOCUSIGN SEND ERROR: %s", e)
        app.logger.error(traceback.format_exc())
        flash("DocuSign send failed: " + str(e))
        if document:
            return redirect(url_for("batch_detail", batch_id=document.batch_id))
        return redirect(request.referrer or url_for("batches_list"))


@app.route("/documents/<int:document_id>/upload-signed", methods=["POST"])
def upload_signed_document(document_id):
    if auth_required():
        return redirect(url_for("login"))

    document = ContractDocument.query.get_or_404(document_id)
    uploaded_file = request.files.get("signed_file")

    if not uploaded_file or not uploaded_file.filename:
        flash("Please choose a signed file to upload.")
        return redirect(url_for("batch_detail", batch_id=document.batch_id))
    if not GOOGLE_DRIVE_FOLDER_ID or not GOOGLE_SERVICE_ACCOUNT_JSON:
        flash("Google Drive is not configured yet.")
        return redirect(url_for("batch_detail", batch_id=document.batch_id))

    file_bytes = uploaded_file.read()
    file_name = uploaded_file.filename
    mime_type = uploaded_file.mimetype or "application/octet-stream"

    try:
        drive_info = upload_bytes_to_drive(
            file_name=file_name,
            file_bytes=file_bytes,
            parent_folder_id=GOOGLE_DRIVE_FOLDER_ID,
            mime_type=mime_type,
        )
    except Exception as e:
        flash("Signed upload failed: " + str(e))
        return redirect(url_for("batch_detail", batch_id=document.batch_id))

    document.signed_file_name = file_name
    document.signed_drive_file_id = drive_info["file_id"]
    document.signed_web_view_link = drive_info["web_view_link"]
    document.signed_uploaded_at = datetime.datetime.utcnow()
    document.status = "signed_uploaded"

    batch_docs = ContractDocument.query.filter_by(batch_id=document.batch_id).all()
    if batch_docs and all(doc.status == "signed_uploaded" for doc in batch_docs):
        document.batch.status = "signed_complete"
    else:
        document.batch.status = "signed_partial"

    db.session.commit()
    flash("Signed file uploaded successfully.")
    return redirect(url_for("batch_detail", batch_id=document.batch_id))


@app.route("/works/<int:work_id>")
def work_detail(work_id):
    if auth_required():
        return redirect(url_for("login"))
    work = Work.query.get_or_404(work_id)
    documents = (
        ContractDocument.query
        .filter_by(work_id=work.id)
        .order_by(ContractDocument.generated_at.desc())
        .all()
    )
    return render_template_string(WORK_DETAIL_HTML, work=work, documents=documents)


try:
    init_db()
except Exception as e:
    print("DB INIT ERROR:", e)

if __name__ == "__main__":
    app.run(debug=True, host="0.0.0.0", port=int(os.getenv("PORT", "5052")))
