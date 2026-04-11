import io
import re
import os
import json
import datetime

from docx import Document
from sqlalchemy import func
from flask import request, session
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload
from docusign_esign import ApiClient

from extensions import db
from config import (
    TEAM_USERNAME, TEAM_PASSWORD,
    DEFAULT_PUBLISHER_ADDRESS, DEFAULT_PUBLISHER_CITY,
    DEFAULT_PUBLISHER_STATE, DEFAULT_PUBLISHER_ZIP,
    GOOGLE_SERVICE_ACCOUNT_JSON, GOOGLE_DRIVE_FOLDER_ID,
    DOCUSIGN_PRIVATE_KEY, DOCUSIGN_INTEGRATION_KEY,
    DOCUSIGN_USER_ID, DOCUSIGN_AUTH_SERVER, DOCUSIGN_BASE_PATH,
)


# ── String helpers ─────────────────────────────────────────────────────────────

def slugify(value):
    value = (value or "").strip()
    value = re.sub(r"[^\w\s-]", "", value, flags=re.UNICODE)
    value = re.sub(r"[-\s]+", "_", value)
    return value or "file"


def parse_float(value):
    try:
        return float((value or "").strip())
    except ValueError:
        return 0.0


def safe_json_loads(value, fallback=None):
    """Parse a JSON string, returning fallback (default []) on any error."""
    if fallback is None:
        fallback = []
    try:
        return json.loads(value) if value else fallback
    except (ValueError, TypeError):
        return fallback


def build_full_name(first_name, middle_name, last_names):
    return " ".join(
        part.strip() for part in [first_name, middle_name, last_names] if part and part.strip()
    ).strip()


def normalize_text(value):
    return " ".join((value or "").lower().strip().split())


def normalize_title(title):
    return normalize_text(title)


# ── Writer identity helpers ────────────────────────────────────────────────────

def build_writer_identity_from_row(row):
    ipi = (row.get("ipi") or "").strip()
    if ipi:
        return "ipi:" + ipi.lower()
    selected_writer_id = (row.get("selected_writer_id") or "").strip()
    if selected_writer_id:
        return "id:" + selected_writer_id
    return "name:" + normalize_text(row.get("full_name", ""))


def build_writer_identity_from_workwriter(work_writer):
    if work_writer.writer and work_writer.writer.ipi:
        return "ipi:" + work_writer.writer.ipi.lower()
    if work_writer.writer_id:
        return "id:" + str(work_writer.writer_id)
    return "name:" + normalize_text(work_writer.writer.full_name if work_writer.writer else "")


# ── Session / publisher helpers ────────────────────────────────────────────────

def build_session_name(raw_name):
    raw_name = (raw_name or "").strip()
    prefix = datetime.datetime.utcnow().strftime("%Y.%m.%d")
    if raw_name:
        return prefix + " " + raw_name
    return prefix


def default_publisher_for_pro(pro):
    return {
        "BMI":   "Songs of Afinarte",
        "ASCAP": "Melodies of Afinarte",
        "SESAC": "Music of Afinarte",
    }.get((pro or "").strip(), "")


def default_publisher_ipi_for_pro(pro):
    return {
        "BMI":   "817874992",
        "ASCAP": "807953316",
        "SESAC": "817094629",
    }.get((pro or "").strip(), "")


# ── External services ──────────────────────────────────────────────────────────

def get_drive_service():
    if not GOOGLE_SERVICE_ACCOUNT_JSON:
        raise ValueError("Missing GOOGLE_SERVICE_ACCOUNT_JSON")
    info = json.loads(GOOGLE_SERVICE_ACCOUNT_JSON)
    credentials = service_account.Credentials.from_service_account_info(
        info, scopes=["https://www.googleapis.com/auth/drive"]
    )
    return build("drive", "v3", credentials=credentials, cache_discovery=False)


def get_docusign_api_client():
    private_key = (DOCUSIGN_PRIVATE_KEY or "").replace("\\n", "\n")
    private_key_bytes = private_key.encode("utf-8")
    api_client = ApiClient()
    api_client.host = DOCUSIGN_BASE_PATH
    token = api_client.request_jwt_user_token(
        client_id=DOCUSIGN_INTEGRATION_KEY,
        user_id=DOCUSIGN_USER_ID,
        oauth_host_name=DOCUSIGN_AUTH_SERVER,
        private_key_bytes=private_key_bytes,
        expires_in=3600,
        scopes=["signature", "impersonation"],
    )
    api_client.set_default_header("Authorization", "Bearer " + token.access_token)
    return api_client


def upload_bytes_to_drive(file_name, file_bytes, parent_folder_id, mime_type):
    service = get_drive_service()
    media = MediaIoBaseUpload(io.BytesIO(file_bytes), mimetype=mime_type, resumable=False)
    metadata = {"name": file_name}
    if parent_folder_id:
        metadata["parents"] = [parent_folder_id]
    created = service.files().create(
        body=metadata,
        media_body=media,
        fields="id, webViewLink",
        supportsAllDrives=True,
    ).execute()
    return {
        "file_id": created.get("id"),
        "web_view_link": created.get("webViewLink"),
    }


# ── List pagination ───────────────────────────────────────────────────────────

class ListPagination:
    """Mimics SQLAlchemy pagination so templates can use the same controls."""
    def __init__(self, items, page, per_page):
        self.total    = len(items)
        self.page     = page
        self.per_page = per_page
        self.pages    = max(1, (self.total + per_page - 1) // per_page)
        self.page     = max(1, min(page, self.pages))
        start         = (self.page - 1) * per_page
        self.items    = items[start:start + per_page]
        self.has_prev = self.page > 1
        self.has_next = self.page < self.pages
        self.prev_num = self.page - 1
        self.next_num = self.page + 1


def paginate_list(items, page, per_page=50):
    return ListPagination(items, page, per_page)


# ── Auth ───────────────────────────────────────────────────────────────────────

FULL_ACCESS_ROLES = {"admin", "label_manager", "publishing_manager"}


def auth_required():
    """Returns True if the request requires login and the user is not authenticated."""
    if session.get("user_id") or session.get("logged_in"):
        return False
    if TEAM_USERNAME and TEAM_PASSWORD:
        return True
    try:
        from models import User
        return User.query.first() is not None
    except Exception:
        return False


def role_required(roles):
    """Returns True if the current user's role is not in the allowed set.
    Call auth_required() first — this only checks the role, not login state."""
    role = session.get("role", "admin")  # legacy team sessions treated as admin
    return role not in roles


# ── Writer helpers ─────────────────────────────────────────────────────────────

def find_existing_writer(selected_writer_id):
    from models import Writer
    if selected_writer_id:
        try:
            writer = Writer.query.get(int(selected_writer_id))
        except (ValueError, TypeError):
            return None
        if writer:
            return writer
    return None


# ── Contract rendering ─────────────────────────────────────────────────────────

def render_docx_template(template_path, data, works_for_table=None):
    if not os.path.exists(template_path):
        raise FileNotFoundError(template_path)
    doc = Document(template_path)

    def replace_all(paragraph):
        text = "".join(run.text for run in paragraph.runs)
        for k, v in data.items():
            text = text.replace("[[" + k + "]]", str(v))
        for run in paragraph.runs:
            run.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = text

    for p in doc.paragraphs:
        replace_all(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_all(p)
    for p in doc.paragraphs:
        if "[[ContractTable]]" in p.text:
            table = doc.add_table(rows=1, cols=5)
            table.style = "Table Grid"
            table.rows[0].cells[0].text = "Work Title"
            table.rows[0].cells[1].text = "Songwriter Name"
            table.rows[0].cells[2].text = "Songwriter Share"
            table.rows[0].cells[3].text = "Publisher Name"
            table.rows[0].cells[4].text = "Publisher Share"
            for item in works_for_table or []:
                row = table.add_row().cells
                row[0].text = item.get("work_title", "")
                row[1].text = item.get("writer_name", "")
                row[2].text = item.get("writer_percentage", "")
                row[3].text = item.get("publisher", "")
                row[4].text = item.get("publisher_percentage", item.get("writer_percentage", ""))
            p.text = ""
            p._element.addnext(table._element)
            break

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ── Form context helpers ───────────────────────────────────────────────────────

def collect_form_context():
    from models import GenerationBatch
    selected_batch_id = (
        request.form.get("existing_batch_id")
        or request.args.get("batch_id")
        or ""
    )
    return {
        "batches": GenerationBatch.query.order_by(GenerationBatch.created_at.desc()).all(),
        "default_publisher_address": DEFAULT_PUBLISHER_ADDRESS,
        "default_publisher_city": DEFAULT_PUBLISHER_CITY,
        "default_publisher_state": DEFAULT_PUBLISHER_STATE,
        "default_publisher_zip": DEFAULT_PUBLISHER_ZIP,
        "force_create": request.form.get("force_create", ""),
        "selected_batch_id": selected_batch_id,
    }


def collect_submitted_form_data():
    return {
        "work_title_value": request.form.get("work_title", ""),
        "contract_date_value": request.form.get("contract_date", ""),
        "new_session_name_value": request.form.get("new_session_name", ""),
        "submitted_writers": [
            {
                "selected_writer_id": writer_id,
                "first_name": first_name,
                "middle_name": middle_name,
                "last_names": last_name,
                "writer_aka": writer_aka,
                "ipi": ipi,
                "email": email,
                "phone_number": phone_number,
                "pro": pro,
                "writer_percentage": percentage,
                "publisher": publisher,
                "publisher_ipi": publisher_ipi,
                "publisher_address": publisher_address,
                "publisher_city": publisher_city,
                "publisher_state": publisher_state,
                "publisher_zip_code": publisher_zip_code,
                "address": address,
                "city": city,
                "state": state,
                "zip_code": zip_code,
            }
            for writer_id, first_name, middle_name, last_name, writer_aka, ipi, email, phone_number, pro, percentage,
                publisher, publisher_ipi, publisher_address, publisher_city, publisher_state, publisher_zip_code,
                address, city, state, zip_code
            in zip(
                request.form.getlist("writer_id"),
                request.form.getlist("writer_first_name"),
                request.form.getlist("writer_middle_name"),
                request.form.getlist("writer_last_names"),
                request.form.getlist("writer_aka"),
                request.form.getlist("writer_ipi"),
                request.form.getlist("writer_email"),
                request.form.getlist("writer_phone_number"),
                request.form.getlist("writer_pro"),
                request.form.getlist("writer_percentage"),
                request.form.getlist("writer_publisher"),
                request.form.getlist("publisher_ipi"),
                request.form.getlist("publisher_address"),
                request.form.getlist("publisher_city"),
                request.form.getlist("publisher_state"),
                request.form.getlist("publisher_zip_code"),
                request.form.getlist("writer_address"),
                request.form.getlist("writer_city"),
                request.form.getlist("writer_state"),
                request.form.getlist("writer_zip_code"),
            )
        ]
    }


def get_batch_writer_summary(batch_id):
    from models import WorkWriter, Work
    work_writers = WorkWriter.query.join(Work).filter(Work.batch_id == batch_id).all()
    grouped = {}
    for ww in work_writers:
        if not ww.writer:
            continue
        if ww.writer_id not in grouped:
            grouped[ww.writer_id] = {"writer": ww.writer, "work_titles": set()}
        grouped[ww.writer_id]["work_titles"].add(ww.work.title)
    summary = []
    for item in grouped.values():
        summary.append({"writer": item["writer"], "work_count": len(item["work_titles"])})
    summary.sort(key=lambda x: x["writer"].full_name.lower())
    return summary


def get_writer_directory_rows(q="", page=1, per_page=50):
    from models import Writer, WorkWriter
    query = Writer.query

    if q:
        like_q = "%" + q.lower() + "%"
        query = query.filter(
            db.or_(
                func.lower(Writer.full_name).like(like_q),
                func.lower(Writer.first_name).like(like_q),
                func.lower(Writer.middle_name).like(like_q),
                func.lower(Writer.last_names).like(like_q),
                func.lower(Writer.writer_aka).like(like_q),
                func.lower(Writer.ipi).like(like_q),
                func.lower(Writer.email).like(like_q),
                func.lower(Writer.phone_number).like(like_q),
            )
        )

    pagination = query.order_by(Writer.full_name.asc()).paginate(page=page, per_page=per_page, error_out=False)
    writers = pagination.items

    writer_ids = [w.id for w in writers]
    work_count_map = dict(
        db.session.query(WorkWriter.writer_id, func.count(WorkWriter.id))
        .filter(WorkWriter.writer_id.in_(writer_ids))
        .group_by(WorkWriter.writer_id)
        .all()
    ) if writer_ids else {}

    rows = [{"writer": w, "work_count": work_count_map.get(w.id, 0)} for w in writers]
    return rows, pagination
